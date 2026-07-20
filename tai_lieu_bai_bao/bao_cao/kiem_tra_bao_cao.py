"""Kiểm tra tệp báo cáo: số liệu, thuật ngữ, cách viết hoa và liên kết giữa các mục.

Năm nhóm kiểm tra:
  1. Mọi con số trong báo cáo phải khớp báo cáo thực nghiệm gốc hoặc tệp kết quả.
  2. Thuật ngữ phải nhất quán: không dùng hai từ khác nhau cho cùng một khái niệm.
  3. Không viết hoa tùy tiện ở giữa câu.
  4. Mục 1 và mục 2 phải liên kết: mọi khái niệm mục 1 nhắc tới đều được mục 2 giải
     thích, và số liệu chung giữa hai mục phải trùng nhau.
  5. Mọi ô số trong mọi bảng phải tìm thấy trong tệp kết quả. Nhóm 1 chỉ soát phần
     chữ, vì python-docx để bảng ngoài danh sách đoạn, nên bảng cần một nhóm riêng.

Chạy: python tai_lieu_bai_bao/bao_cao/kiem_tra_bao_cao.py
"""

from __future__ import annotations

import os
import re
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
BAO_CAO = os.path.join(HERE, "Bao_cao_dong_gop_PIE-Net.docx")
GOC = (r"C:\Users\trong\AppData\Local\Temp\claude"
       r"\c--Users-trong-Downloads-PIE-em-Hoang"
       r"\2bf96163-73da-4142-8ed0-60cb16987721\scratchpad\bao_cao_goc.txt")

# Số liệu và nguồn phải đối chiếu. Khóa là con số như viết trong báo cáo.
NGUON_SO = {
    "28,31": "goc", "28,91": "goc", "0,60": "goc", "0,991": "goc",
    "26,80": "goc", "25,79": "goc", "26,82": "goc", "24,64": "goc",
    "0,59": "goc", "0,88": "goc", "1,97": "goc", "27,47": "goc",
    "27,60": "goc", "8,2": "goc", "7,6": "goc", "0,3": "goc", "0,1": "goc",
    "13,2": "ket_qua", "19,2": "ket_qua", "4,7": "ket_qua", "9,7": "ket_qua",
    "17,2": "ket_qua", "23,6": "ket_qua", "7,2": "ket_qua", "14,7": "ket_qua",
    "1,0000": "ket_qua",
    # số của phép đo nội bộ, đối chiếu với tài liệu trong tai_lieu_bai_bao
    "24,3458": "noi_bo", "1,242": "noi_bo", "1,243": "noi_bo",
    # mức phần dư biến phân mục tiêu, ấn định trước, có trong tệp kết quả
    "3,0": "ket_qua", "1,0": "ket_qua",
    # bảng 6.1 và cái giá của tiêu chuẩn kiểm được, từ certificate_check.csv
    # và certificate_cost.csv
    "5,29": "ket_qua", "2,83": "ket_qua", "2,18": "ket_qua", "11,96": "ket_qua",
    "0,0901": "ket_qua", "0,0110": "ket_qua", "2,6": "ket_qua", "5,8": "ket_qua",
    "3,10": "ket_qua", "1,62": "ket_qua", "1,63": "ket_qua", "0,819": "ket_qua",
    "1,06": "ket_qua", "0,462": "ket_qua", "0,643": "ket_qua", "0,226": "ket_qua",
    "0,332": "ket_qua", "0,0872": "ket_qua", "0,218": "ket_qua", "0,0505": "ket_qua",
    "2,26": "ket_qua", "2,52": "ket_qua", "3,03": "ket_qua", "3,98": "ket_qua",
    "5,41": "ket_qua", "6,95": "ket_qua",
    # mục 7: hệ số chi phí, mức khả thi, chi phí của lịch quá siết, độ dốc tổng được
    "563,77": "ket_qua", "476,69": "ket_qua", "732,89": "ket_qua",
    "1,2433": "ket_qua", "42,83": "ket_qua", "39,64": "ket_qua",
    "2,342": "ket_qua", "2,838": "ket_qua", "1,191": "ket_qua", "1,011": "ket_qua",
    "1,006": "ket_qua", "1,000": "ket_qua",
    # mục 8: hai con số trước khi sửa cách đo, và bẫy lịch bước tăng tốc
    "12,86": "ket_qua", "14,90": "ket_qua", "2,37": "ket_qua", "3,46": "ket_qua",
    # tham số do ta ấn định trước khi chạy, không phải kết quả đo
    "0,9": "thiet_ke", "1,01": "thiet_ke",
    # số suy ra từ một số khác đã đối chiếu: 24,3 phần trăm là từ mức khả thi 1,2433
    "24,3": "suy_ra",
    # mức khả thi của hai cấu hình dừng hẳn bên trong tập ràng buộc, và hai tham số
    # thực nghiệm ấn định trước
    "0,9994": "ket_qua", "0,9995": "ket_qua",
    "0,05": "thiet_ke", "0,55": "thiet_ke",
}

# Tài liệu nội bộ dùng làm nguồn đối chiếu cho các phép đo của chính đề tài.
NOI_BO = [os.path.join(HERE, "..", f) for f in
          ("04_lo_trinh_chung_minh.md", "05_chung_minh_hoi_tu_yeu.md",
           "02_ket_qua_so_tien_hoa.md")]
KET_QUA_DIR = os.path.join(HERE, "..", "..", "results", "theory")

# Cặp thuật ngữ: (từ đúng, các từ đồng nghĩa không được dùng).
THUAT_NGU = [
    ("phép chiếu xấp xỉ", ["chiếu gần đúng", "phép chiếu không chính xác",
                           "chiếu xấp xỉ hoá"]),
    ("phần dư biến phân", ["sai số hội tụ", "độ dư", "residual"]),
    ("bước nội", ["vòng trong", "lặp trong"]),
    ("bước ngoài", ["vòng ngoài", "lặp ngoài"]),
    ("khởi tạo ấm", ["warm start", "khởi động ấm"]),
    ("tập ràng buộc", ["miền ràng buộc", "tập khả thi"]),
    ("chứng chỉ", ["giấy chứng nhận", "certificate"]),
    ("hệ số vô hướng", ["hệ số scalar", "trọng số vô hướng"]),
]

# Từ được phép viết hoa giữa câu: tên riêng, tên phương pháp và ký hiệu toán học.
CHO_PHEP_HOA = {
    "Bx", "B", "D", "F", "x", "y",          # ký hiệu toán học trong y = Bx + ε
    "Numerical", "Algorithms", "Journal", "Scientific", "Computing",
    "Optimization", "Computational", "Applications", "Communications",
    "Nonlinear", "Science", "Simulation", "SIAM", "Ferreira", "Ugon",
    "Millán", "Qin", "Tan", "Tseng",        # tên tạp chí và tên tác giả
    "Plug-and-Play", "Gauss", "PIE-Net", "Fejér", "DnCNN", "Q1", "Q2",
    "Chambolle-Pock", "Malitsky", "Lipschitz", "Hà", "Nội", "Đào", "Trọng", "Hiếu",
    "Đặng", "Văn", "Chiến", "Học", "Viện", "Bộ", "Khoa", "Công", "Nghệ",
    "Bưu", "Chính", "Viễn", "Thông", "Mục", "Đề", "Giảng", "Lớp", "Biết",
    "Báo", "Tóm", "Bối", "Câu", "Khẳng", "Kết", "Điểm", "Đóng", "Phần",
    "Về", "Từ", "Có", "Cách", "Thiết", "Vì", "Điều", "Trên", "Đây", "Ba",
    "Hai", "Thứ", "Hệ", "Ngoài", "Nội", "Với", "Nhờ", "Số", "Trước",
}


BO_QUA_SO = {"4,1"}          # số hiệu bảng, không phải số liệu


def khop_lam_tron(so_bao_cao, nguon):
    """Số trong báo cáo có thể là số trong nguồn đã làm tròn.

    Ví dụ báo cáo ghi 19,2 còn tệp kết quả ghi 19,17. So khớp chuỗi cứng sẽ báo sai;
    ở đây ta tìm trong nguồn một số mà khi làm tròn tới đúng số chữ số thập phân của
    số trong báo cáo thì trùng với nó. Nguồn có thể ghi số ở dạng khoa học, chẳng hạn
    9.012900e-02, nên phải nhận cả dạng đó."""
    muc = so_bao_cao.replace(",", ".")
    if muc in nguon:
        return True
    try:
        gia_tri = float(muc)
    except ValueError:
        return False
    so_le = len(muc.split(".")[1]) if "." in muc else 0
    for ung_vien in re.findall(r"\d+\.\d+(?:[eE][+-]?\d+)?", nguon):
        if round(float(ung_vien), so_le) == gia_tri:
            return True
    return False


def lay_van_ban(path):
    d = Document(path)
    return [p.text.strip() for p in d.paragraphs if p.text.strip()]


def lay_o_bang(path):
    """Trả về các ô của mọi bảng, kèm vị trí để báo lỗi cho rõ."""
    d = Document(path)
    ra = []
    for bi, t in enumerate(d.tables, start=1):
        for ri, hang in enumerate(t.rows):
            for ci, o in enumerate(hang.cells):
                ra.append((bi, ri, ci, o.text.strip()))
    return ra


def kiem_bang(o_bang, ket_qua):
    """Mọi ô chứa số trong bảng phải truy được về tệp kết quả.

    Bỏ qua ô chỉ là số nguyên nhỏ dùng làm nhãn, và ô đã khai báo trong NGUON_SO
    với nguồn không phải kết quả đo."""
    loi = []
    for bi, ri, ci, chu in o_bang:
        for so in re.findall(r"\d+,\d+", chu):
            if NGUON_SO.get(so) in ("goc", "noi_bo", "thiet_ke", "suy_ra"):
                continue
            if not khop_lam_tron(so, ket_qua):
                loi.append(f"bảng {bi}, dòng {ri}, cột {ci}: số {so} "
                           f"không tìm thấy trong tệp kết quả")
    return loi


def tach_muc(doan):
    """Tách danh sách đoạn thành từng mục, trả về dict {số mục: các đoạn}."""
    # Đòi có khoảng trắng ngay sau dấu chấm, để câu mở đầu bằng số tiểu mục như
    # "Mục 5.3 đã hoãn lại..." không bị nhận nhầm là tiêu đề của mục 5.
    moc = [(i, int(t.split(".")[0].split()[-1]))
           for i, t in enumerate(doan) if re.match(r"^Mục \d+\. \S", t)]
    ket = {}
    for k, (i, so) in enumerate(moc):
        j = moc[k + 1][0] if k + 1 < len(moc) else len(doan)
        ket[so] = doan[i:j]
    return ket


def kiem_so_lieu(text, goc, noi_bo, ket_qua):
    """Mọi con số trong báo cáo phải tìm thấy ở nguồn tương ứng.

    Cảnh báo cả trường hợp ngược lại: con số xuất hiện trong báo cáo mà chưa có
    trong danh sách đối chiếu, tức là một con số chưa được kiểm."""
    loi = []
    for so, nguon in NGUON_SO.items():
        if so not in text:
            continue
        if nguon == "goc" and so not in goc:
            loi.append(f"số {so} không tìm thấy trong báo cáo gốc")
        if nguon == "noi_bo" and so not in noi_bo:
            loi.append(f"số {so} không tìm thấy trong tài liệu nội bộ")
        if nguon == "ket_qua" and not khop_lam_tron(so, ket_qua):
            loi.append(f"số {so} không tìm thấy trong tệp kết quả, "
                       f"kể cả khi tính tới làm tròn")
    for so in sorted(set(re.findall(r"\d+,\d+", text))):
        if so not in NGUON_SO and so not in BO_QUA_SO:
            loi.append(f"số {so} chưa được đưa vào danh sách đối chiếu")
    return loi


def kiem_thuat_ngu(doan):
    """Soát TỪNG đoạn riêng, không nối các đoạn lại.

    Nếu nối, chữ cuối của đoạn này và chữ đầu của đoạn sau sẽ tạo ra những cụm từ
    không có thật, chẳng hạn tiêu đề kết thúc bằng 'trùng lặp' đứng ngay trước đoạn
    mở đầu bằng 'Trong' sẽ bị đọc thành 'lặp trong'."""
    loi = []
    for t in doan:
        thap = t.lower()
        for dung, sai_list in THUAT_NGU:
            for sai in sai_list:
                if sai.lower() in thap:
                    loi.append(f"dùng '{sai}' thay vì '{dung}'")
    return loi


def kiem_viet_hoa(doan):
    loi = []
    for t in doan:
        # bỏ qua tiêu đề mục và các dòng bìa
        if re.match(r"^(Mục|\d\.\d\.)", t) or len(t) < 60:
            continue
        for cau in re.split(r"(?<=\.)\s+", t):
            tu = cau.split()
            for w in tu[1:]:
                w_sach = w.strip(".,;:()")
                if (w_sach and w_sach[0].isupper() and w_sach not in CHO_PHEP_HOA
                        and not w_sach.isupper()):
                    loi.append(f"viết hoa giữa câu: '{w_sach}'")
    return loi


def kiem_lien_ket(muc):
    """Kiểm liên kết theo chuỗi giữa các mục đã viết.

    Ba tiêu chí: mỗi mục phải dẫn sang mục kế tiếp; khái niệm nêu ở mục trước phải
    được mục sau nói tới; và số liệu dùng chung giữa các mục phải trùng nhau."""
    canh_bao = []
    so_muc = sorted(muc)
    van = {k: " ".join(v).lower() for k, v in muc.items()}

    # a) mỗi mục dẫn sang mục kế tiếp
    for k in so_muc[:-1]:
        if f"mục {k + 1}" not in van[k]:
            canh_bao.append(f"mục {k} không có câu dẫn sang mục {k + 1}")

    # b) khái niệm nêu ở mục 1 phải được mục 2 giải thích
    for k in ["bốn khẳng định", "plug-and-play", "khởi tạo ấm",
              "phép chiếu xấp xỉ", "cấu trúc"]:
        if 1 in van and 2 in van and k in van[1] and k not in van[2]:
            canh_bao.append(f"mục 1 nêu '{k}' nhưng mục 2 không nhắc lại")

    # c) ba lỗi phương pháp luận nêu ở mục 1 phải được mục 3 trình bày
    if 1 in van and 3 in van and "ba lỗi" in van[1]:
        for k in ["thời gian", "mục tiêu", "bất đối xứng"]:
            if k not in van[3]:
                canh_bao.append(f"mục 1 nhắc ba lỗi nhưng mục 3 thiếu '{k}'")

    # d) số liệu dùng chung giữa các cặp mục
    for i, a in enumerate(so_muc):
        for b in so_muc[i + 1:]:
            chung = set(re.findall(r"\d+,\d+", van[a])) & \
                    set(re.findall(r"\d+,\d+", van[b]))
            if chung:
                canh_bao.append(
                    f"(thông tin) số dùng chung mục {a} và {b}: {sorted(chung)}")
    return canh_bao


def main():
    if not os.path.exists(BAO_CAO):
        sys.exit(f"Không tìm thấy báo cáo: {BAO_CAO}")
    doan = lay_van_ban(BAO_CAO)
    o_bang = lay_o_bang(BAO_CAO)
    goc = open(GOC, encoding="utf-8").read() if os.path.exists(GOC) else ""
    noi_bo = "".join(open(f, encoding="utf-8").read()
                     for f in NOI_BO if os.path.exists(f))
    ket_qua = ""
    if os.path.isdir(KET_QUA_DIR):
        for f in os.listdir(KET_QUA_DIR):
            if f.endswith(".csv") or f.endswith(".log"):
                ket_qua += open(os.path.join(KET_QUA_DIR, f),
                                encoding="utf-8").read()
    muc = tach_muc(doan)
    text = " ".join(doan)

    print(f"Báo cáo: {len(doan)} đoạn, {len(text.split())} chữ, "
          f"{len(o_bang)} ô bảng")
    for k in sorted(muc):
        sc = len(" ".join(muc[k]).split())
        print(f"   mục {k}: {len(muc[k]):2d} đoạn, {sc:4d} chữ "
              f"(khoảng {sc / 550:.1f} trang)")
    print()

    tat_ca_loi = []
    for ten, loi in [
        ("1. Đối chiếu số liệu với nguồn", kiem_so_lieu(text, goc, noi_bo, ket_qua)),
        ("2. Thuật ngữ nhất quán", kiem_thuat_ngu(doan)),
        ("3. Viết hoa giữa câu", kiem_viet_hoa(doan)),
        ("5. Số trong bảng", kiem_bang(o_bang, ket_qua)),
    ]:
        print(f"--- {ten} ---")
        if loi:
            for l in sorted(set(loi)):
                print("   LỖI:", l)
            tat_ca_loi += loi
        else:
            print("   không có lỗi")
        print()

    print("--- 4. Liên kết giữa các mục ---")
    for c in kiem_lien_ket(muc):
        print("  ", c)
        if not c.startswith("(thông tin)"):
            tat_ca_loi.append(c)
    print()

    print("=" * 60)
    print("KẾT QUẢ:", "đạt, không có lỗi" if not tat_ca_loi
          else f"có {len(set(tat_ca_loi))} lỗi cần sửa")


if __name__ == "__main__":
    main()
