"""Sinh tệp báo cáo, gồm bìa và mục 1.

Cách làm: dùng tệp báo cáo SoundGuard làm khuôn cho trang bìa, để giữ nguyên khung
viền hoa văn và logo học viện (hai ảnh này nằm ở đoạn 0 và đoạn 3 của tệp gốc, không
phải là đường kẻ nên không thể tạo lại bằng thiết lập viền trang). Chương trình chỉ
thay phần chữ trên bìa, xóa toàn bộ nội dung cũ phía sau, rồi ghi mục 1 vào.

Chạy: python tai_lieu_bai_bao/bao_cao/tao_bao_cao.py
"""

from __future__ import annotations

import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
KHUON = r"C:\Users\trong\Downloads\SoundGuard_Report_FULL_EN.docx"
OUT = os.path.join(HERE, "Bao_cao_dong_gop_PIE-Net.docx")

# Chữ trên bìa, theo đúng thứ tự đoạn của tệp khuôn.
# Khóa là chỉ số đoạn; đoạn 0 và 3 chứa ảnh nên không đụng tới.
BIA = {
    1: "BỘ KHOA HỌC VÀ CÔNG NGHỆ",
    2: "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
    4: "Biết khi nào phép chiếu đã đủ chính xác",
    5: ("Đề tài: chứng chỉ sai số tính được cho phép chiếu xấp xỉ và phân tích chi "
        "phí trong phương pháp chiếu phản xạ giải bất đẳng thức biến phân (PIE-Net)"),
    6: "Giảng viên hướng dẫn: Đặng Văn Chiến",
    7: "Học viên: Đào Trọng Hiếu — M25VMCS07",
    8: "Lớp: M25VMCS",
    9: "Hà Nội - 2026",
}

DOAN_NGAT_TRANG = 10          # đoạn chứa ngắt trang, giữ lại
MUC_1 = [
    ("Mục 1. Tóm tắt", "de_muc"),
    ("Báo cáo này trình bày các đóng góp của giai đoạn nghiên cứu tiếp nối báo cáo "
     "thực nghiệm trước đó của đề tài PIE-Net. Điểm xuất phát là một kết luận không "
     "thuận lợi. Báo cáo thực nghiệm trước đã kiểm bốn khẳng định của thiết kế ban "
     "đầu qua năm thí nghiệm, mỗi thí nghiệm có tiêu chí đạt hay không đạt đặt trước "
     "khi chạy, và chỉ một khẳng định đứng vững. Hệ số vô hướng học được đạt 28,31 dB "
     "so với 28,91 dB của hệ số hằng được tinh chỉnh tốt, tức thấp hơn 0,60 dB; phương "
     "pháp của đề tài thua Plug-and-Play 0,59 dB ở chế độ khớp mức nhiễu và 0,88 dB ở "
     "chế độ lệch mức nhiễu. Các kết quả này có tính cấu trúc, tức xuất phát từ bản "
     "chất toán học của cách dựng mô hình chứ không từ lỗi cài đặt, nên hướng thực "
     "nghiệm đã được đóng lại. Khẳng định duy nhất còn đứng vững là lợi thế chi phí "
     "của phép chiếu xấp xỉ có khởi tạo ấm.", "thuong"),
    ("Câu hỏi của giai đoạn này vì thế được thu hẹp lại: có thể biến điều duy nhất "
     "còn đứng vững đó thành một đóng góp công bố được hay không. Để trả lời, trọng "
     "tâm được chuyển từ phần học sang phần thuật toán và giải tích số, sơ đồ lặp "
     "được dựng lại cho đúng với tài liệu, và toàn bộ phần đo chi phí được thiết kế "
     "lại.", "thuong"),
    ("Báo cáo có ba đóng góp. Đóng góp thứ nhất là một chứng chỉ sai số tính được cho "
     "phép chiếu xấp xỉ. Định nghĩa phép chiếu xấp xỉ trong tài liệu đòi kiểm khoảng "
     "cách tới phép chiếu chính xác, mà phép chiếu chính xác lại chính là đại lượng "
     "thuật toán được thiết kế để tránh tính; đây là một vòng luẩn quẩn khiến chế độ "
     "ngân sách trong thực thi nằm ngoài phạm vi của định lý. Vì bài toán chiếu có hàm "
     "mục tiêu lồi mạnh, khoảng cách tới nghiệm chiếu bị chặn bởi căn của hai lần "
     "khoảng cách đối ngẫu, và khoảng cách đối ngẫu tính trực tiếp từ cặp biến gốc và "
     "đối ngẫu mà bộ giải nội đã có sẵn. Nhờ đó tiêu chuẩn dừng trở thành thứ một "
     "chương trình có thể kiểm được.", "thuong"),
    ("Đóng góp thứ hai là chế độ ngân sách thích nghi dựa trên chứng chỉ đó, cùng với "
     "kết quả chi phí. Trên bài toán khử mờ ảnh với tập ràng buộc là quả cầu biến phân "
     "toàn phần, để đạt cùng một mức phần dư biến phân, chế độ thích nghi rẻ hơn phép "
     "chiếu chính xác từ 13,2 đến 19,2 lần tính theo tổng bước nội và từ 4,7 đến 9,7 "
     "lần tính theo thời gian thuật toán trên mờ Gauss; trên mờ chuyển động các hệ số "
     "tương ứng là 17,2 đến 23,6 lần và 7,2 đến 14,7 lần. Mọi cấu hình đều cho nghiệm "
     "khả thi, với mức vi phạm ràng buộc không vượt 1,0000, trong khi chế độ ngân sách "
     "cố định cho đầu ra vi phạm ràng buộc tới vài phần trăm.", "thuong"),
    ("Đóng góp thứ ba là một giao thức đo chi phí công bằng, được xây dựng sau khi ba "
     "lỗi phương pháp luận trong chính phần đo của chúng tôi bị phát hiện và sửa: chỉ "
     "đếm bước nội mà không đo thời gian, trong khi bước nội của chế độ thích nghi đắt "
     "hơn vì phải tính chứng chỉ; lấy mức phần dư mục tiêu từ chính kết quả của phương "
     "pháp đối chứng, tức chọn sau khi đã thấy số liệu; và dò tham số bất đối xứng, "
     "mười sáu cấu hình cho phương pháp đề xuất so với hai cấu hình cho phương pháp "
     "đối chứng. Giao thức sau khi sửa báo cả hai thước đo chi phí, ấn định mức phần "
     "dư mục tiêu trước khi chạy, và dò tham số cho hai bên với cùng số cấu hình.",
     "thuong"),
    ("Về mức công bố, báo cáo kết luận rằng công trình phù hợp một tạp chí thuộc nhóm "
     "Q2 chứ không phải nhóm Q1. Lý do nằm ở phần lý thuyết: sau khi dẫn xuất đầy đủ, "
     "hằng số của số hạng nhiễu do phép chiếu xấp xỉ không chứa nghịch đảo độ dài bước, "
     "nên việc thêm bước phản xạ vào khung phép chiếu xấp xỉ không tạo ra một cơ chế "
     "nhiễu mới về bản chất; định lý hội tụ vì thế là một mở rộng của các kết quả đã "
     "công bố. Giá trị công bố được của công trình nằm ở phần chứng chỉ, chế độ ngân "
     "sách thích nghi và phân tích chi phí, tức ở phần thuật toán và thực nghiệm.",
     "thuong"),
    ("Phần chứng minh chưa hoàn tất ở hai chi tiết kỹ thuật: phát biểu chính xác bổ đề "
     "tựa Fejér loại nhân tính sao cho nó không hấp thụ mất số hạng âm, và bước chuyển "
     "giới hạn yếu trong không gian vô hạn chiều; chi tiết thứ hai là thường quy trong "
     "không gian hữu hạn chiều của phần thực nghiệm.", "thuong"),
    ("Phần còn lại của báo cáo được tổ chức như sau. Mục 2 trình bày bối cảnh và điểm "
     "xuất phát, tức bài toán và các kết luận của báo cáo thực nghiệm trước. Mục 3 nêu "
     "bốn nguyên tắc làm việc, là cơ sở để đánh giá độ tin cậy của các con số về sau. "
     "Các mục từ 4 đến 9 trình bày lần lượt các đóng góp. Mục 10 ghi lại các kết quả "
     "âm tính, mục 11 mô tả chất lượng mã nguồn và tính tái lập, và các mục cuối nêu "
     "hạn chế cùng việc còn lại.", "thuong"),
]

MUC_2 = [
    ("Mục 2. Bối cảnh và điểm xuất phát", "de_muc"),

    ("Mục 1 đã liệt kê các đóng góp nhưng chưa nói vì sao chúng cần thiết. Mục này trả "
     "lời câu hỏi đó: trình bày bài toán, nhắc lại các kết luận của báo cáo thực nghiệm "
     "trước, và chỉ ra điều gì còn đứng vững sau khi phần lớn các kết luận ấy bị đóng "
     "lại. Chính phần còn lại đó là điểm xuất phát của giai đoạn nghiên cứu này.",
     "thuong"),

    ("2.1. Bài toán", "de_muc_phu"),
    ("Đề tài PIE-Net nhằm giải bài toán ngược trong xử lý ảnh, viết dưới dạng y = Bx + "
     "ε, trong đó x là ảnh gốc cần khôi phục, y là dữ liệu quan sát, B là toán tử suy "
     "biến đã biết, chẳng hạn một phép làm mờ, và ε là nhiễu. Bài toán này được đặt "
     "dưới dạng bất đẳng thức biến phân trên một tập ràng buộc D: tìm điểm x thuộc D "
     "sao cho tích vô hướng của toán tử chi phí tại x với mọi hướng đi từ x tới một "
     "điểm khác của D đều không âm.", "thuong"),
    ("Cách đặt bài toán như vậy có hai hệ quả định hình toàn bộ công việc về sau. Hệ "
     "quả thứ nhất là tập ràng buộc trở thành một thành phần của thuật toán chứ không "
     "phải một điều kiện phụ, vì mỗi bước lặp đều phải chiếu điểm hiện tại về tập ràng "
     "buộc. Hệ quả thứ hai, và đây là điểm quan trọng cho phần sau của báo cáo, là "
     "phép chiếu chỉ rẻ khi tập ràng buộc đủ đơn giản. Với tập ràng buộc được dùng "
     "trong đề tài là quả cầu biến phân toàn phần, phép chiếu không có công thức đóng "
     "và phải giải bằng một vòng lặp nội, nên chi phí của phép chiếu chiếm phần lớn "
     "chi phí của thuật toán.", "thuong"),

    ("2.2. Kết luận của báo cáo thực nghiệm trước", "de_muc_phu"),
    ("Thiết kế ban đầu của đề tài dựa trên bốn khẳng định, mỗi khẳng định được gắn với "
     "một thí nghiệm kiểm chứng và một tiêu chí đạt hay không đạt bằng con số cụ thể, "
     "đặt trước khi chạy và không sửa sau khi thấy số liệu. Báo cáo thực nghiệm trước "
     "đã kiểm cả bốn khẳng định qua năm thí nghiệm, và kết quả là chỉ một khẳng định "
     "đứng vững.", "thuong"),
    ("Khẳng định thứ nhất cho rằng hệ số vô hướng học được cải thiện chất lượng khôi "
     "phục so với hệ số hằng được tinh chỉnh tốt. Kết quả là không đạt: phiên bản có "
     "hệ số vô hướng học được đạt 28,31 dB, còn phiên bản hệ số hằng đạt 28,91 dB, tức "
     "thấp hơn 0,60 dB. Hệ số học được còn tự hội tụ về một hằng số xấp xỉ 1, với giá "
     "trị trung bình 0,991 trên tập kiểm tra, cho thấy phần dung lượng mô hình dành "
     "cho nó không được sử dụng.", "thuong"),
    ("Khẳng định thứ ba cho rằng thành phần học được thêm giá trị so với phiên bản "
     "không học. Kết quả cũng không đạt: tại điểm vận hành ổn định nhất, với mờ Gauss "
     "phiên bản có học đạt 26,80 dB còn phiên bản không học đạt 25,79 dB, và với mờ "
     "chuyển động là 26,82 dB so với 24,64 dB; sau khi tính tới dao động theo hạt "
     "giống, mức chênh không vượt ngưỡng 0,3 dB đặt trước một cách ổn định. Ở ngân "
     "sách bước nội thấp, phiên bản có học còn mất ổn định nghiêm trọng.", "thuong"),
    ("Khẳng định thứ tư cho rằng ràng buộc cứng nhất quán dữ liệu cho lợi thế so với "
     "Plug-and-Play. Kết quả là không đạt, và đây là kết quả dứt khoát nhất: phương "
     "pháp của đề tài thua Plug-and-Play 0,59 dB ở chế độ khớp mức nhiễu và 0,88 dB ở "
     "chế độ lệch mức nhiễu, tức khoảng cách còn nới rộng đúng ở chế độ mà ràng buộc "
     "cứng được kỳ vọng có lợi thế. Trước đó một thí nghiệm riêng đã đo được khoảng "
     "trống 1,97 dB giữa prior thủ công tốt nhất và trần của việc học trên ảnh kết "
     "cấu, nhưng chính thí nghiệm đối đầu cho thấy Plug-and-Play đã lấy gần trọn "
     "khoảng trống đó, khi đạt 27,47 dB so với trần 27,60 dB.", "thuong"),

    ("2.3. Vì sao các kết quả này có tính cấu trúc", "de_muc_phu"),
    ("Điểm quan trọng nhất của báo cáo trước không nằm ở các con số mà ở chẩn đoán "
     "nguyên nhân. Ba kết quả không đạt nêu trên không đến từ lỗi cài đặt hay từ việc "
     "tinh chỉnh chưa tới, mà bắt nguồn từ chính cách dựng mô hình, nên không thể khắc "
     "phục bằng cách dò thêm siêu tham số hay đổi kiến trúc mạng.", "thuong"),
    ("Có thể thấy điều đó qua ba lập luận. Thứ nhất, hệ số vô hướng dương nhân với một "
     "toán tử không làm thay đổi tập nghiệm của bất đẳng thức biến phân; nó chỉ tác "
     "động lên động học hội tụ, mà ở chân trời hữu hạn thì tác động đó quá nhỏ để tạo "
     "khác biệt về chất lượng. Thứ hai, ràng buộc cứng nhất quán dữ liệu buộc nghiệm "
     "bám sát quả cầu quanh dữ liệu quan sát, mà dữ liệu quan sát đã mang nhiễu, nên "
     "ép như vậy chính là nạp nhiễu trở lại ảnh khôi phục; điều này giải thích vì sao "
     "khoảng cách với Plug-and-Play nới rộng đúng ở chế độ lệch mức nhiễu. Thứ ba, hai "
     "thành phần được kỳ vọng của đề tài loại trừ lẫn nhau: toán tử học được chỉ ổn "
     "định khi ngân sách bước nội đủ lớn, mà ngân sách bước nội lớn lại phá hủy đúng "
     "lợi thế chi phí của phép chiếu xấp xỉ ít bước nội.", "thuong"),
    ("Vì các rào cản này có tính cấu trúc, báo cáo trước đã khuyến nghị đóng hướng "
     "thực nghiệm theo khuôn khổ cũ thay vì tiếp tục tinh chỉnh. Báo cáo hiện tại kế "
     "thừa nguyên vẹn khuyến nghị đó và không mở lại bất kỳ nhánh nào đã đóng.",
     "thuong"),

    ("2.4. Điều duy nhất còn đứng vững", "de_muc_phu"),
    ("Khẳng định thứ hai là khẳng định duy nhất đạt tiêu chí đặt trước. Nội dung của "
     "nó là: trên một tập ràng buộc mà phép chiếu chính xác thật sự cần vòng lặp nội, "
     "phép chiếu xấp xỉ với ngân sách bước nội nhỏ đạt chất lượng không thấp hơn phép "
     "chiếu chính xác có khởi tạo ấm, trong ngưỡng 0,1 dB, đồng thời tốn ít hơn rõ rệt "
     "về tổng bước nội. Số liệu cụ thể là rẻ hơn khoảng 8,2 lần với mờ Gauss và khoảng "
     "7,6 lần với mờ chuyển động.", "thuong"),
    ("Báo cáo trước cũng nêu ba điều kiện để hiểu đúng kết quả này, và báo cáo hiện "
     "tại giữ nguyên cả ba. Thứ nhất, lợi thế đo được là lợi thế về chi phí chứ không "
     "phải về chất lượng khôi phục. Thứ hai, phép chiếu chính xác có khởi tạo ấm mới "
     "là mốc so sánh đúng, vì bản thân khởi tạo ấm đã tiết kiệm rất lớn so với khởi "
     "tạo lạnh, và nếu so với khởi tạo lạnh thì phần tiết kiệm ấy sẽ bị gán nhầm cho "
     "phép chiếu xấp xỉ. Thứ ba, cơ chế tạo ra lợi thế là việc phân bổ vòng lặp nội "
     "qua các bước ngoài nhờ khởi tạo ấm, một kỹ thuật đã biết trong tối ưu, nên tự nó "
     "chưa đủ mới để đứng thành một đóng góp riêng.", "thuong"),

    ("2.5. Câu hỏi của giai đoạn nghiên cứu này", "de_muc_phu"),
    ("Từ hiện trạng trên, câu hỏi của giai đoạn nghiên cứu này được phát biểu như sau: "
     "có thể biến điều duy nhất còn đứng vững, tức lợi thế chi phí của phép chiếu xấp "
     "xỉ có khởi tạo ấm, thành một đóng góp công bố được hay không, khi mà bản thân cơ "
     "chế tạo ra lợi thế đó là một kỹ thuật đã biết.", "thuong"),
    ("Câu hỏi này ràng buộc cách làm theo ba hướng. Thứ nhất, không được lặp lại các "
     "nhánh đã đóng, nên trọng tâm phải chuyển từ phần học sang phần thuật toán và "
     "giải tích số. Thứ hai, vì cơ chế nền đã biết, đóng góp nếu có phải nằm ở chỗ "
     "khác, và phần sau của báo cáo cho thấy chỗ đó là tiêu chuẩn dừng của vòng lặp "
     "nội. Thứ ba, vì kết luận trước đây dựa trên so sánh chi phí, mọi phép đo chi phí "
     "trong giai đoạn này phải được thiết kế chặt hơn trước, và mục 3 trình bày các "
     "nguyên tắc được dùng cho việc đó.", "thuong"),
]

MUC_3 = [
    ("Mục 3. Phương pháp làm việc", "de_muc"),

    ("Mục 2 kết thúc bằng một yêu cầu: giai đoạn nghiên cứu này phải được thiết kế chặt "
     "hơn giai đoạn trước, để không lặp lại việc rút kết luận từ bằng chứng yếu. "
     "Mục này trình bày bốn nguyên tắc được áp dụng thống nhất cho toàn bộ giai đoạn "
     "nghiên cứu. Lý do phải trình bày chúng trước khi nêu kết quả là như sau: đóng "
     "góp của báo cáo này chủ yếu là một so sánh chi phí, mà một so sánh chi phí chỉ "
     "có nghĩa khi phép đo công bằng, phương pháp đối chứng được cho cơ hội mạnh nhất, "
     "và tiêu chí không bị điều chỉnh sau khi đã thấy số liệu. Bốn nguyên tắc dưới đây "
     "nhằm bảo đảm đúng ba điều đó.", "thuong"),

    ("3.1. Nguyên tắc thứ nhất: đặt tiêu chí trước khi chạy", "de_muc_phu"),
    ("Mọi so sánh chi phí trong báo cáo đều được thực hiện tại các mức phần dư biến "
     "phân mục tiêu ấn định trước khi chạy, cụ thể là năm mức từ 3,0 nhân mười mũ trừ "
     "hai xuống 1,0 nhân mười mũ trừ hai. Các mức này độc lập với kết quả của mọi cấu "
     "hình, và không được sửa sau khi có số liệu.", "thuong"),
    ("Nguyên tắc này không phải hình thức. Bản đo đầu tiên của chúng tôi lấy mức phần "
     "dư mục tiêu từ chính phần dư cuối của phương pháp đối chứng rồi nới thêm năm "
     "phần trăm. Cách làm đó tưởng là tự nhiên nhưng thực chất là chọn tiêu chí sau "
     "khi đã thấy số liệu, và nó có lợi một cách cơ học cho phương pháp được đề xuất. "
     "Lỗi này được phát hiện ở một vòng phản biện và đã được sửa bằng cách chuyển sang "
     "danh sách mức ấn định trước.", "thuong"),

    ("3.2. Nguyên tắc thứ hai: cho phương pháp đối chứng cơ hội mạnh nhất", "de_muc_phu"),
    ("Một kết luận rằng phương pháp này rẻ hơn phương pháp kia chỉ có giá trị khi "
     "phương pháp đối chứng đã được chỉnh tốt. Vì vậy phương pháp đối chứng, tức phép "
     "chiếu chính xác dừng theo chứng chỉ, được dò tham số với đúng số cấu hình như "
     "phương pháp được đề xuất, cụ thể là tám cấu hình cho mỗi bên.", "thuong"),
    ("Nguyên tắc này cũng xuất phát từ một lỗi thật. Bản đo trước đó dò mười sáu cấu "
     "hình cho phương pháp được đề xuất nhưng chỉ hai cấu hình cho phương pháp đối "
     "chứng. Sự bất đối xứng ấy làm hệ số tiết kiệm cao hơn thực tế, vì phương pháp "
     "được đề xuất được chọn ở điểm vận hành tốt nhất còn phương pháp đối chứng thì "
     "không.", "thuong"),
    ("Cùng nguyên tắc này, báo cáo giữ lại mốc so sánh khó nhất chứ không chọn mốc dễ. "
     "Phép chiếu chính xác luôn được cho khởi tạo ấm, vì bản thân khởi tạo ấm đã tiết "
     "kiệm rất lớn; nếu so với phép chiếu chính xác khởi tạo lạnh thì phần tiết kiệm "
     "đó sẽ bị gán nhầm cho phép chiếu xấp xỉ. Ngoài ra, hai bên đều dừng theo cùng "
     "một chứng chỉ tính được, nên không bên nào được dùng thông tin mà bên kia không "
     "có.", "thuong"),

    ("3.3. Nguyên tắc thứ ba: phản biện đối kháng trước khi chấp nhận", "de_muc_phu"),
    ("Mỗi bản thảo chứng minh và mỗi kết luận số quan trọng đều được đưa qua một vòng "
     "phản biện độc lập trước khi được chấp nhận. Người phản biện được yêu cầu tự tính "
     "lại từng phép biến đổi và tìm cách bác bỏ kết luận, chứ không phải xác nhận nó.",
     "thuong"),
    ("Kết quả của nguyên tắc này là hai bản thảo chứng minh đã bị bác và ba lỗi phương "
     "pháp luận trong chính phần đo của chúng tôi đã bị phát hiện. Bản thảo chứng minh "
     "thứ nhất bị bác vì một lỗ hổng ở phần lõi: bổ đề một bước được mượn đòi hỏi một "
     "chuỗi phép chiếu liền mạch, mà bước quán tính và bước neo trong sơ đồ khi đó làm "
     "đứt chuỗi ấy. Bản thảo thứ hai bị bác vì một lập luận về tính mới không đứng "
     "vững. Ba lỗi phương pháp luận gồm: chỉ đếm bước nội mà không đo thời gian; chọn "
     "mức phần dư mục tiêu sau khi thấy số liệu; và dò tham số bất đối xứng giữa hai "
     "bên.", "thuong"),
    ("Cần nói rõ rằng việc liệt kê các lần bị bác không phải là tự phê bình cho có. "
     "Trong bốn lỗi nghiêm trọng nhất của giai đoạn này, không lỗi nào được phát hiện "
     "bởi tác giả khi đang viết, mà tất cả đều do vòng phản biện tìm ra. Điều này cho "
     "thấy vòng phản biện là bộ phận không thể bỏ của quy trình chứ không phải một "
     "bước kiểm tra thêm.", "thuong"),

    ("3.4. Nguyên tắc thứ tư: ghi lại kết quả âm tính", "de_muc_phu"),
    ("Mọi hướng đã thử và thất bại đều được ghi lại kèm số liệu và kèm chẩn đoán "
     "nguyên nhân, thay vì bị xóa đi. Trong giai đoạn này có bốn hướng như vậy: lịch "
     "bước tăng tốc cho bài toán chiếu; siết chứng chỉ bằng cách theo dõi giá trị tốt "
     "nhất; tiêu chuẩn sai số tương đối theo chuẩn toán tử; và tiêu chuẩn chiếu nới "
     "lỏng. Mục 10 trình bày đầy đủ bốn hướng này.", "thuong"),
    ("Việc ghi lại có hai lợi ích cụ thể. Lợi ích thứ nhất là tránh lặp lại: mỗi hướng "
     "âm tính đều kèm lý do thất bại đủ rõ để người sau không thử lại, ví dụ tiêu "
     "chuẩn chiếu nới lỏng không thực thi được vì việc kiểm nó đòi tính hàm tựa của "
     "tập ràng buộc, tức đắt ngang chính phép chiếu. Lợi ích thứ hai là các kết quả âm "
     "tính giúp xác định đúng phạm vi của kết quả dương tính, vì chúng cho biết phương "
     "pháp hỏng ở đâu chứ không chỉ chạy tốt ở đâu.", "thuong"),

    ("3.5. Hệ quả đối với cách đọc báo cáo", "de_muc_phu"),
    ("Bốn nguyên tắc trên dẫn tới một hệ quả về cách đọc các con số ở những mục sau. "
     "Các hệ số tiết kiệm được báo cáo là hệ số đo sau khi đã sửa ba lỗi phương pháp "
     "luận nêu trên, nên chúng thấp hơn các con số từng xuất hiện trong những bản đo "
     "trước đó của chính chúng tôi. Đây là điều có chủ ý: bản đo cuối cùng chặt hơn "
     "các bản trước, và con số thấp hơn là con số đáng tin hơn.", "thuong"),
    ("Ngoài ra, mỗi con số trong báo cáo đều được đối chiếu tự động với tệp kết quả "
     "gốc bằng một chương trình kiểm tra, thay vì chép tay. Quy trình đối chiếu này đã "
     "phát hiện hai con số sai trong bản thảo đầu của chính báo cáo này, gồm một hệ số "
     "lấy nhầm từ bản đo cũ và một mức vi phạm ràng buộc ghi không chính xác. Mục 11 "
     "mô tả chương trình kiểm tra đó.", "thuong"),
    ("Mục 4 sau đây áp dụng ngay các nguyên tắc trên vào bước đầu tiên của giai đoạn "
     "nghiên cứu, là việc chọn hướng đi.", "thuong"),
]

# Bảng ở mục 4: mỗi dòng gồm bài chặn và điểm khác biệt so với công việc của báo cáo.
BANG_MUC_4 = {
    "tieu_de": "Bảng 4.1. Các bài chặn và điểm khác biệt so với công việc của báo cáo.",
    "cot": ["Bài chặn", "Điểm khác biệt"],
    "dong": [
        ("Malitsky 2015, phương pháp chiếu phản xạ, SIAM Journal on Optimization",
         "Dùng bước phản xạ nhưng với phép chiếu chính xác; không xét phép chiếu "
         "xấp xỉ."),
        ("Díaz Millán, Ferreira và Ugon 2024, Computational Optimization and "
         "Applications",
         "Dùng phép chiếu xấp xỉ với tiêu chuẩn sai số tương đối, nhưng trên sơ đồ "
         "ngoại suy hai phép chiếu mỗi bước, không có bước phản xạ."),
        ("Bài về phép chiếu xấp xỉ có độ nhớt 2025, Communications in Nonlinear "
         "Science and Numerical Simulation",
         "Ghép phép chiếu xấp xỉ với bước độ nhớt, nhưng không có bước phản xạ và "
         "không phân tích chi phí của vòng lặp nội."),
        ("Tan và Qin 2020",
         "Có quán tính và độ nhớt với một phép chiếu mỗi bước, nhưng dùng hiệu chỉnh "
         "kiểu Tseng và phép chiếu chính xác."),
        ("Bản cải tiến của phương pháp chiếu phản xạ 2023, Numerical Algorithms",
         "Tổng quát hóa quy tắc chọn độ dài bước, không phải phép chiếu xấp xỉ; đã kiểm "
         "riêng vì đăng trên chính tạp chí được nhắm tới."),
        ("Bản có quán tính của phương pháp chiếu phản xạ 2022, Journal of Scientific "
         "Computing",
         "Đã ghép bước phản xạ với quán tính, nên phần quán tính không còn là chỗ "
         "trống."),
    ],
}

MUC_4 = [
    ("Mục 4. Định vị lại hướng nghiên cứu", "de_muc"),

    ("Sau khi hướng thực nghiệm cũ được đóng lại, việc đầu tiên phải làm là chọn hướng "
     "đi mới. Mục này trình bày cách chọn hướng đó. Nguyên tắc áp dụng ở đây là nguyên "
     "tắc thứ nhất của mục 3: quyết định phải dựa trên bằng chứng đặt ra trước, chứ "
     "không dựa vào cảm nhận rằng một hướng nào đó có vẻ mới.", "thuong"),

    ("4.1. Cách khảo sát tài liệu", "de_muc_phu"),
    ("Khảo sát được tổ chức thành năm góc độc lập, mỗi góc phụ trách một câu hỏi riêng "
     "và trả về một danh sách bài kèm các thuộc tính phân biệt. Năm góc gồm: phép "
     "chiếu xấp xỉ cho bất đẳng thức biến phân; các phương pháp có quán tính và độ "
     "nhớt cho lớp toán tử giả đơn điệu; các phương pháp kiểu neo và bước phản xạ; "
     "điều kiện để có tốc độ hội tụ; và các phương pháp học được có bảo đảm hội tụ cho "
     "bài toán ngược trong xử lý ảnh.", "thuong"),
    ("Với mỗi bài tìm được, khảo sát ghi lại sáu thuộc tính: lớp toán tử được giả "
     "thiết; có bước quán tính hay không; có bước độ nhớt hoặc bước neo hay không; "
     "dùng phép chiếu chính xác hay phép chiếu xấp xỉ; nếu dùng phép chiếu xấp xỉ thì "
     "tiêu chuẩn sai số là gì; và có phân tích chi phí của vòng lặp nội hay không. "
     "Cách ghi theo thuộc tính như vậy cho phép so sánh trực tiếp giữa các bài, thay "
     "vì chỉ đọc phần tóm tắt rồi kết luận cảm tính.", "thuong"),

    ("4.2. Kết quả khảo sát: tài liệu chia thành hai nhánh", "de_muc_phu"),
    ("Kết quả khảo sát cho thấy tài liệu hiện có chia thành hai nhánh gần như không "
     "giao nhau. Nhánh thứ nhất gồm các phương pháp có bước quán tính, bước độ nhớt "
     "hoặc bước neo, được nghiên cứu rất kỹ và cho hội tụ mạnh trong nhiều trường hợp; "
     "nhưng toàn bộ nhánh này giả thiết phép chiếu là chính xác. Nhánh thứ hai gồm các "
     "phương pháp dùng phép chiếu xấp xỉ, có xét tới sai số của phép chiếu; nhưng các "
     "bài trong nhánh này không có bước quán tính và không có bước độ nhớt, và quan "
     "trọng hơn, không bài nào phân tích chi phí thực của vòng lặp nội.", "thuong"),
    ("Sự phân đôi này là điểm mấu chốt của khảo sát, vì nó chỉ ra chỗ trống: không có bài "
     "nào đồng thời dùng bước phản xạ và phép chiếu xấp xỉ. Bảng 4.1 liệt kê các bài "
     "chặn quan trọng nhất cùng điểm khác biệt của từng bài so với công việc của báo "
     "cáo này.", "thuong"),
    ("__BANG_4_1__", "bang"),

    ("4.3. Khe hở còn lại, và việc thừa nhận nó hẹp", "de_muc_phu"),
    ("Từ bảng trên, khe hở còn lại được phát biểu chính xác như sau: ghép bước phản xạ "
     "với phép chiếu xấp xỉ. Bước phản xạ đáng chú ý về mặt chi phí vì nó chỉ cần tính "
     "toán tử một lần mỗi bước ngoài, trong khi các sơ đồ ngoại suy thường cần hai "
     "lần.", "thuong"),
    ("Cần thừa nhận ngay rằng đây là một khe hở hẹp. Mọi thành phần của sơ đồ đều đã "
     "có người dùng riêng lẻ: bước phản xạ đã có từ 2015, phép chiếu xấp xỉ với tiêu "
     "chuẩn sai số tương đối đã có từ 2024, và việc ghép bước phản xạ với quán tính đã "
     "có từ 2022. Phần chưa ai làm chỉ là đúng cặp ghép còn lại. Báo cáo này không "
     "trình bày khe hở đó như một phát hiện lớn, và mục 9 sẽ cho thấy phần định lý thu "
     "được từ khe hở này chỉ là một mở rộng.", "thuong"),
    ("Chính vì khe hở hẹp, hướng nghiên cứu được chuyển trọng tâm ngay từ đầu: thay vì "
     "cố tìm tính mới ở phần định lý, công việc tập trung vào một câu hỏi thực thi mà "
     "cả hai nhánh tài liệu đều bỏ ngỏ, đó là làm sao kiểm được tiêu chuẩn dừng của "
     "vòng lặp nội. Mục 6 trình bày câu trả lời cho câu hỏi đó.", "thuong"),

    ("4.4. Kiểm rủi ro trùng lặp", "de_muc_phu"),
    ("Trong khảo sát có một bài đáng lo hơn cả, là bản cải tiến của phương pháp chiếu "
     "phản xạ đăng năm 2023 trên Numerical Algorithms, tức chính tạp chí được nhắm "
     "tới. Tên bài chứa chữ tổng quát hóa, mà trong dòng nghiên cứu này chữ đó thường "
     "hàm ý đã bao gồm cả nhiễu hoặc phép chiếu xấp xỉ. Nếu đúng như vậy thì toàn bộ "
     "khe hở nêu trên không còn, và mọi công việc sau đó là vô ích.", "thuong"),
    ("Vì rủi ro này có thể làm hỏng cả hướng nghiên cứu, nó được kiểm trước mọi việc "
     "khác. Kết quả kiểm cho thấy bài đó tổng quát hóa quy tắc chọn độ dài bước chứ không "
     "phải phép chiếu; nó cho hội tụ yếu và tốc độ tuyến tính theo nghĩa R khi toán tử "
     "đơn điệu mạnh, và không xét phép chiếu xấp xỉ. Khe hở vì thế vẫn còn.", "thuong"),
    ("Cùng lần kiểm này, khảo sát phát hiện thêm một bài năm 2022 trên Journal of "
     "Scientific Computing đã ghép bước phản xạ với quán tính. Phát hiện đó không phá "
     "khe hở, nhưng nó loại bỏ một khả năng: phần quán tính không còn là chỗ trống, "
     "nên không nên tìm đóng góp ở đó. Đây là một trong các căn cứ để về sau bỏ hẳn "
     "bước quán tính khỏi sơ đồ, như mục 5 trình bày.", "thuong"),
]

MUC_5 = [
    ("Mục 5. Sửa sai trong thiết kế và danh pháp", "de_muc"),

    ("Mục này trình bày bốn việc sửa sai đã thực hiện trước khi bắt đầu phần đóng góp "
     "chính. Cả bốn đều không hào nhoáng, và ba trong số đó là sửa lỗi của chính đề "
     "tài. Tuy vậy chúng được trình bày thành một mục riêng vì hai lý do: nếu không "
     "sửa thì mọi kết quả sau đó đứng trên nền sai, và bản thân cách phát hiện ra "
     "chúng minh họa cho các nguyên tắc làm việc ở mục 3.", "thuong"),

    ("5.1. Sửa sai về danh pháp", "de_muc_phu"),
    ("Sơ đồ lặp ban đầu của đề tài gồm bốn pha, trong đó pha thứ ba được gọi là hiệu "
     "chỉnh phản xạ. Cách gọi này sai. Pha đó có dạng lấy điểm chiếu rồi trừ đi hiệu "
     "của toán tử tại hai điểm, tức là hiệu chỉnh kiểu Tseng, và nó tính toán tử hai "
     "lần mỗi bước ngoài. Bước phản xạ kiểu Malitsky là một thứ khác hẳn: nó dùng điểm "
     "phản xạ, tức điểm đối xứng lấy từ hai bước lặp gần nhất, làm đối số của toán tử, "
     "và chỉ tính toán tử một lần mỗi bước ngoài.", "thuong"),
    ("Hệ quả của việc gọi sai tên không nằm ở chữ nghĩa mà ở việc định vị. Khi còn gọi "
     "pha thứ ba là hiệu chỉnh phản xạ, sơ đồ trông như một phương pháp có bước phản "
     "xạ, tức thuộc một nhánh tài liệu tương đối trống. Khi gọi đúng tên là hiệu chỉnh "
     "kiểu Tseng, sơ đồ hiện ra gần trùng với một bài đã công bố năm 2020, là bài Tan "
     "và Qin đã nêu ở bảng 4.1, vốn có đủ quán tính, độ nhớt, một phép chiếu mỗi bước "
     "và hiệu chỉnh kiểu Tseng. Nói cách khác, cách gọi sai đã che mờ một sự trùng lặp "
     "với tài liệu, và chỉ sau khi sửa tên thì mới thấy phải đổi sơ đồ chứ không phải "
     "chỉ đổi cách trình bày.", "thuong"),
    ("Việc sửa này dẫn tới một quyết định cụ thể: chuyển pha thứ ba từ hiệu chỉnh kiểu "
     "Tseng sang bước phản xạ kiểu Malitsky đúng nghĩa. Ngoài lý do định vị, việc "
     "chuyển còn có lợi về chi phí, vì bước phản xạ chỉ cần tính toán tử một lần mỗi "
     "bước ngoài.", "thuong"),

    ("5.2. Bỏ bước quán tính", "de_muc_phu"),
    ("Bước quán tính được bỏ khỏi sơ đồ. Có hai lý do độc lập, một về lý thuyết và một "
     "về số liệu, và điều đáng nói là chúng dẫn tới cùng một kết luận.", "thuong"),
    ("Lý do về lý thuyết là bước quán tính làm đứt chuỗi phép chiếu mà bổ đề một bước "
     "đòi hỏi. Bổ đề một bước của phương pháp chiếu phản xạ được xây trên giả thiết "
     "rằng điểm gốc của bước sau chính là kết quả phép chiếu của bước trước, tạo thành "
     "một chuỗi liền mạch. Bước quán tính chèn một điểm ngoại suy vào giữa hai bước "
     "chiếu, nên chuỗi ấy bị cắt và bổ đề không dùng được. Đây chính là lỗ hổng đã làm "
     "một bản thảo chứng minh bị bác, như đã nêu ở mục 3.", "thuong"),
    ("Lý do về số liệu là bỏ bước quán tính không làm mất kết quả. Phép đo trực tiếp "
     "trên mờ Gauss, với ngân sách hai bước nội và 150 bước ngoài, cho thấy phiên bản "
     "có bước quán tính đạt 24,3458 dB với phần dư biến phân 1,242 nhân mười mũ trừ "
     "hai, còn phiên bản không có bước quán tính đạt 24,3458 dB với phần dư 1,243 nhân "
     "mười mũ trừ hai. Hai kết quả trùng khít về chất lượng khôi phục và chênh nhau "
     "khoảng một phần nghìn về phần dư biến phân.", "thuong"),
    ("Kết hợp với phát hiện ở mục 4.4 rằng việc ghép bước phản xạ với quán tính đã có "
     "người làm từ năm 2022, quyết định bỏ bước quán tính là rõ ràng: nó không đóng "
     "góp gì về số liệu, nó cản trở chứng minh, và nếu giữ thì phần đó cũng không còn "
     "là chỗ trống trong tài liệu.", "thuong"),

    ("5.3. Bỏ bước neo", "de_muc_phu"),
    ("Bước neo, tức bước kéo dãy lặp về một điểm cố định để có hội tụ mạnh, cũng được "
     "bỏ. Cái giá phải trả là kết quả hội tụ thu được chỉ là hội tụ yếu thay vì hội tụ "
     "mạnh; đây là điều báo cáo thừa nhận chứ không né tránh. Đổi lại có hai lý do, và "
     "cũng như trường hợp bước quán tính, chúng độc lập nhau.", "thuong"),
    ("Lý do thứ nhất giống lý do ở mục 5.2: bước neo cũng chèn một phép trộn vào giữa "
     "hai bước chiếu, nên cũng làm đứt chuỗi phép chiếu. Khi bỏ cả bước quán tính lẫn "
     "bước neo, điểm lặp tiếp theo chính là kết quả phép chiếu, chuỗi liền mạch được "
     "khôi phục, và bổ đề một bước dùng được. Cần nói thêm rằng với toán tử đơn điệu, "
     "hội tụ yếu là kết quả chuẩn của dòng nghiên cứu này, nên việc chấp nhận hội tụ "
     "yếu không phải là hạ tiêu chuẩn.", "thuong"),
    ("Lý do thứ hai chỉ lộ ra sau một phép đo dài và có phần bất ngờ, nên được trình "
     "bày đầy đủ ở mục 7. Tóm tắt ở đây: bước neo làm cho dãy sai số của phép chiếu "
     "xấp xỉ nằm sát ranh giới không tổng được, tức sát ranh giới vi phạm chính giả "
     "thiết mà định lý cần. Như vậy bỏ bước neo không chỉ giúp chứng minh chạy được mà "
     "còn cứu một giả thiết của chính định lý đó.", "thuong"),

    ("5.4. Sửa một niềm tin sai về ràng buộc hộp", "de_muc_phu"),
    ("Phát hiện thứ tư khác ba phát hiện trên ở chỗ nó sửa một niềm tin sai đã tồn tại "
     "lâu trong đề tài, chứ không sửa một thành phần của sơ đồ. Niềm tin đó là: phép "
     "kẹp giá trị điểm ảnh về ràng buộc hộp làm hỏng cấu trúc toán tử gần kề của bài "
     "toán chiếu, nên phải bỏ ràng buộc hộp khi chạy phần lý thuyết. Niềm tin này xuất "
     "hiện trong cả tài liệu nội bộ lẫn một vòng phản biện trước đó, nên nó được coi "
     "là đúng suốt một thời gian dài.", "thuong"),
    ("Niềm tin đó sai. Toán tử gần kề của tổng nửa chuẩn bình phương với hàm chỉ của "
     "hộp đúng bằng phép kẹp áp lên nghiệm của phần bậc hai, vì ràng buộc hộp tách "
     "được theo từng tọa độ nên bài toán rã thành các bài toán một chiều độc lập. Điều "
     "này được kiểm bằng tìm kiếm vét cạn trên lưới, và sai lệch giữa công thức kẹp "
     "với nghiệm tìm được chỉ ở mức một phần mười nghìn, đúng bằng độ mịn của lưới.",
     "thuong"),
    ("Hệ quả của việc sửa niềm tin này lớn hơn vẻ ngoài của nó. Ràng buộc hộp là ràng "
     "buộc vật lý tự nhiên của ảnh, vì giá trị điểm ảnh vốn nằm trong một khoảng đóng. "
     "Quan trọng hơn, quả cầu biến phân toàn phần tự nó không bị chặn, còn giao của nó "
     "với hộp thì bị chặn; tính bị chặn ấy là điều kiện cần cho một hướng đã được thử "
     "ở mục 10. Nếu vẫn giữ niềm tin sai và loại bỏ ràng buộc hộp, hướng đó thậm chí "
     "không thể thử được.", "thuong"),

    ("5.5. Sơ đồ sau khi sửa", "de_muc_phu"),
    ("Sau bốn việc sửa trên, sơ đồ lặp rút từ bốn pha xuống còn hai dòng: lấy điểm "
     "phản xạ từ hai bước lặp gần nhất, rồi chiếu xấp xỉ điểm thu được sau một bước đi "
     "theo hướng ngược với toán tử tại điểm phản xạ. Toán tử chỉ được tính một lần mỗi "
     "bước ngoài, và chuỗi phép chiếu liền mạch được giữ.", "thuong"),
    ("Việc rút gọn này không phải là làm nghèo sơ đồ mà là loại bỏ những thành phần "
     "không đóng góp. Ba thành phần bị bỏ đều đã được kiểm riêng: bước quán tính không "
     "đổi kết quả về số liệu, bước neo đổi lại một giả thiết nằm sát ranh giới, và "
     "hiệu chỉnh kiểu Tseng làm sơ đồ trùng với tài liệu. Sơ đồ hai dòng này là sơ đồ "
     "được dùng cho toàn bộ phần còn lại của báo cáo, và mục 6 bắt đầu từ đây để trình "
     "bày đóng góp chính.", "thuong"),
]

# Bảng ở mục 6: kiểm chứng chặn trên của chứng chỉ theo số bước nội.
BANG_MUC_6 = {
    "tieu_de": ("Bảng 6.1. Kiểm chứng chặn trên của chứng chỉ theo số bước nội. "
                "Cột cuối là tỉ số giữa chặn trên và sai số thật."),
    "cot": ["Số bước nội", "Chặn trên từ chứng chỉ", "Sai số thật",
            "Tỉ số chặn trên trên sai số thật"],
    "dong": [
        ("5", "5,29", "2,83", "2,18"),
        ("10", "3,10", "1,62", "2,26"),
        ("25", "1,63", "0,819", "2,52"),
        ("50", "1,06", "0,462", "3,03"),
        ("100", "0,643", "0,226", "3,98"),
        ("250", "0,332", "0,0872", "5,41"),
        ("500", "0,218", "0,0505", "6,95"),
        ("1500", "0,0901", "0,0110", "11,96"),
    ],
}

MUC_6 = [
    ("Mục 6. Chứng chỉ sai số tính được", "de_muc"),

    ("Mục này trình bày đóng góp kỹ thuật chính của giai đoạn nghiên cứu. Nội dung của "
     "nó là gỡ một vòng luẩn quẩn nằm ngay trong định nghĩa của phép chiếu xấp xỉ, và "
     "nhờ đó nối được phần lý thuyết với phần thực thi. Mục 4 đã chỉ ra rằng cả hai "
     "nhánh tài liệu đều bỏ ngỏ câu hỏi này, nên đây cũng là chỗ đáng đầu tư nhất.",
     "thuong"),

    ("6.1. Vấn đề: một vòng luẩn quẩn trong định nghĩa", "de_muc_phu"),
    ("Trong tài liệu, phép chiếu xấp xỉ với sai số cho trước được định nghĩa như sau: "
     "một điểm được gọi là phép chiếu xấp xỉ của một điểm cho trước nếu nó nằm trong "
     "tập ràng buộc và khoảng cách từ nó tới phép chiếu chính xác không vượt quá mức "
     "sai số cho phép. Định nghĩa này rõ ràng về mặt toán học và là định nghĩa đúng để "
     "phát biểu định lý.", "thuong"),
    ("Nhưng nó không kiểm được trong thực thi. Muốn biết một điểm có thỏa định nghĩa "
     "hay không, chương trình phải tính khoảng cách từ điểm đó tới phép chiếu chính "
     "xác, tức phải biết phép chiếu chính xác. Trong khi đó, toàn bộ lý do dùng phép "
     "chiếu xấp xỉ là để tránh phải tính phép chiếu chính xác, vì trên quả cầu biến "
     "phân toàn phần phép chiếu chính xác đòi một vòng lặp nội tốn kém. Nói cách khác, "
     "tiêu chuẩn dừng đòi hỏi đúng đại lượng mà thuật toán sinh ra để tránh.", "thuong"),
    ("Hệ quả của vòng luẩn quẩn này không chỉ là bất tiện. Vì tiêu chuẩn không kiểm "
     "được, các cách chọn ngân sách bước nội trong thực thi, chẳng hạn chạy một số "
     "bước nội cố định, đều nằm ngoài phạm vi của định lý: chúng có thể chạy tốt trên "
     "số liệu nhưng không ai bảo đảm chúng thỏa giả thiết mà định lý cần. Đây là "
     "khoảng cách giữa lý thuyết và thực thi mà mục này nhằm xóa bỏ.", "thuong"),

    ("6.2. Cách gỡ: dùng khoảng cách đối ngẫu làm chứng chỉ", "de_muc_phu"),
    ("Cách gỡ dựa trên một tính chất của chính bài toán chiếu. Phép chiếu một điểm lên "
     "tập ràng buộc là nghiệm của bài toán cực tiểu hóa nửa bình phương khoảng cách "
     "tới điểm đó trên tập ràng buộc. Hàm mục tiêu của bài toán này lồi mạnh với tham "
     "số bằng một, vì phần bậc hai của nó chính là nửa bình phương khoảng cách.",
     "thuong"),
    ("Với một hàm lồi mạnh tham số bằng một, khoảng cách bình phương từ một điểm bất "
     "kỳ tới nghiệm không vượt quá hai lần độ chênh giữa giá trị hàm mục tiêu tại điểm "
     "đó và giá trị tối ưu. Độ chênh ấy lại không vượt quá khoảng cách đối ngẫu, tức "
     "hiệu giữa giá trị hàm mục tiêu gốc tại điểm đang xét và giá trị hàm mục tiêu đối "
     "ngẫu tại biến đối ngẫu đang xét. Ghép hai bước lại, khoảng cách từ điểm đang xét "
     "tới phép chiếu chính xác không vượt quá căn bậc hai của hai lần khoảng cách đối "
     "ngẫu.", "thuong"),
    ("Điểm mấu chốt là khoảng cách đối ngẫu tính được. Bộ giải nội được dùng là thuật "
     "toán Chambolle-Pock, vốn là thuật toán gốc và đối ngẫu, nên ở mỗi bước nội nó "
     "sẵn có cả biến gốc lẫn biến đối ngẫu. Chỉ cần thay hai biến đó vào hai hàm mục "
     "tiêu là ra khoảng cách đối ngẫu, mà không cần biết nghiệm. Vòng luẩn quẩn vì thế "
     "được gỡ: tiêu chuẩn dừng trở thành thứ chương trình kiểm được bằng chính thông "
     "tin nó đang có.", "thuong"),

    ("6.3. Công thức cho bài toán chiếu lên quả cầu biến phân toàn phần", "de_muc_phu"),
    ("Với tập ràng buộc là quả cầu biến phân toàn phần, bài toán chiếu được viết dưới "
     "dạng cực tiểu hóa tổng của hai hàm: hàm thứ nhất là nửa bình phương khoảng cách "
     "tới điểm cần chiếu, hàm thứ hai là hàm chỉ của quả cầu chuẩn hỗn hợp áp lên "
     "trường gradient rời rạc. Hàm mục tiêu gốc, tính tại một điểm đã được ép về tập "
     "ràng buộc, chính là nửa bình phương khoảng cách từ điểm đó tới điểm cần chiếu.",
     "thuong"),
    ("Hàm mục tiêu đối ngẫu, tính tại biến đối ngẫu, gồm ba số hạng: tích vô hướng "
     "giữa điểm cần chiếu với ảnh của biến đối ngẫu qua toán tử phân kỳ lấy dấu âm; "
     "trừ đi nửa bình phương chuẩn của ảnh đó; và trừ đi tích của bán kính quả cầu với "
     "chuẩn hỗn hợp cực đại của biến đối ngẫu. Khoảng cách đối ngẫu là hiệu giữa hàm "
     "mục tiêu gốc và hàm mục tiêu đối ngẫu, và nó luôn không âm.", "thuong"),
    ("Trong cài đặt, điểm được ép về tập ràng buộc bằng cách co về giá trị trung bình "
     "của chính nó. Phép co này giữ được tính khả thi vì biến phân toàn phần của một "
     "ảnh hằng bằng không và biến phân toàn phần thuần nhất dương theo phép co, nên "
     "chỉ cần co với hệ số bằng tỉ số giữa bán kính và biến phân toàn phần hiện tại "
     "khi giá trị hiện tại vượt bán kính, còn khi đã khả thi thì giữ nguyên. "
     "Nhờ vậy điểm trả về vừa khả thi vừa có chứng chỉ đi kèm.", "thuong"),

    ("6.4. Vị thế của đóng góp", "de_muc_phu"),
    ("Cần nêu rõ vị thế của đóng góp này để tránh hiểu quá lời. Bất đẳng thức nối "
     "khoảng cách tới nghiệm với khoảng cách đối ngẫu dưới giả thiết lồi mạnh là kiến "
     "thức chuẩn trong giải tích lồi; nó không phải phát hiện của báo cáo, và một vòng "
     "phản biện đã nhắc đúng điều đó.", "thuong"),
    ("Đóng góp nằm ở chỗ khác: dùng bất đẳng thức ấy làm giao diện giữa phần lý thuyết "
     "và bộ giải nội. Phần lý thuyết phát biểu điều kiện theo mức sai số cho phép, còn "
     "bộ giải nội chỉ đo được các đại lượng của chính nó; chứng chỉ là thứ dịch giữa "
     "hai bên. Các bài về phép chiếu xấp xỉ đều bỏ ngỏ chỗ này: hoặc giả thiết dãy sai "
     "số một cách trừu tượng, hoặc dùng tiêu chuẩn cần biết phép chiếu chính xác. Việc "
     "lấp chỗ ấy là điều mà báo cáo này nhận là đóng góp, không hơn.", "thuong"),

    ("6.5. Kiểm chứng: chứng chỉ có thật sự là chặn trên hay không", "de_muc_phu"),
    ("Một chứng chỉ chỉ dùng được nếu nó không bao giờ đánh giá thấp sai số thật. Nếu "
     "nó đánh giá thấp, thuật toán sẽ dừng sớm và giả thiết của định lý bị vi phạm mà "
     "không ai biết. Vì vậy tính chất chặn trên được kiểm trực tiếp bằng số, chứ không "
     "chỉ dựa vào lập luận.", "thuong"),
    ("Phép kiểm được thực hiện như sau: chạy bộ giải nội tới nhiều mức số bước khác "
     "nhau; ở mỗi mức, tính chặn trên từ chứng chỉ, đồng thời tính sai số thật bằng "
     "cách so với một nghiệm chiếu tham chiếu chạy rất dài; rồi so hai đại lượng. Kết "
     "quả trình bày ở bảng 6.1.", "thuong"),
    ("__BANG_6_1__", "bang6"),
    ("Ở toàn bộ tám mức số bước nội đã thử, chặn trên luôn lớn hơn sai số thật, nên "
     "chứng chỉ hợp lệ. Ngoài ra cả hai đại lượng cùng giảm khi số bước nội tăng, đúng "
     "như mong đợi: ở năm bước nội chặn trên là 5,29 còn sai số thật là 2,83; ở 1500 "
     "bước nội chặn trên còn 0,0901 và sai số thật còn 0,0110.", "thuong"),

    ("6.6. Cái giá phải trả", "de_muc_phu"),
    ("Chứng chỉ có một nhược điểm phải nói thẳng: nó bi quan, và mức bi quan tăng dần "
     "khi tiến gần nghiệm. Cột cuối của bảng 6.1 cho thấy tỉ số giữa chặn trên và sai "
     "số thật tăng từ 2,18 ở năm bước nội lên 11,96 ở 1500 bước nội. Nói cách khác, "
     "càng gần nghiệm thì chứng chỉ càng đánh giá sai số cao hơn thực tế nhiều lần.",
     "thuong"),
    ("Hệ quả trực tiếp là chi phí. Dừng theo chứng chỉ tốn nhiều bước nội hơn dừng "
     "theo sai số thật từ 2,6 đến 5,8 lần, tùy mức sai số yêu cầu. Đây là cái giá của "
     "một tiêu chuẩn kiểm được: sai số thật là đại lượng không biết được trong thực "
     "thi, nên con số 2,6 đến 5,8 lần chính là khoảng cách giữa một tiêu chuẩn lý "
     "tưởng nhưng không dùng được và một tiêu chuẩn dùng được.", "thuong"),
    ("Có một cách đã được thử để giảm mức bi quan, là theo dõi giá trị tốt nhất của "
     "hai hàm mục tiêu qua các bước nội thay vì chỉ dùng giá trị ở bước hiện tại. Cách "
     "này về lý thuyết cho chặn trên không lỏng hơn, nhưng đo thực tế cho thấy nó "
     "không giảm được bước nội nào, vì thuật toán Chambolle-Pock vốn đã cho giá trị "
     "gần đơn điệu. Kết quả âm tính này được ghi ở mục 10.", "thuong"),
    ("Dù bi quan, chứng chỉ vẫn đáng dùng, vì nó biến một tiêu chuẩn không kiểm được "
     "thành một tiêu chuẩn kiểm được, và vì nó mở đường cho chế độ ngân sách thích "
     "nghi trình bày ở mục 7.", "thuong"),
]


# Hai bảng ở mục 7: chi phí của chế độ thích nghi so với phép chiếu chính xác,
# trích từ results/theory/muc7_he_so_gauss.csv và muc7_he_so_motion.csv.
COT_MUC_7 = ["Mức phần dư biến phân", "Bước nội của phép chiếu chính xác",
             "Giây của phép chiếu chính xác", "Bước nội của chế độ thích nghi",
             "Giây của chế độ thích nghi", "Hệ số theo bước nội",
             "Hệ số theo thời gian"]

BANG_MUC_7A = {
    "tieu_de": ("Bảng 7.1. Chi phí trên mờ Gauss ở năm mức phần dư biến phân ấn định "
                "trước. Hai cột cuối là số lần chế độ thích nghi rẻ hơn."),
    "cot": COT_MUC_7,
    "dong": [
        ("3,0·10⁻²", "3796", "2,48", "198", "0,52", "19,17", "4,73"),
        ("2,0·10⁻²", "3858", "5,13", "275", "0,70", "14,03", "7,38"),
        ("1,5·10⁻²", "5738", "8,61", "357", "0,88", "16,07", "9,74"),
        ("1,2·10⁻²", "8509", "9,73", "476", "1,12", "17,88", "8,68"),
        ("1,0·10⁻²", "8566", "11,25", "649", "1,43", "13,20", "7,89"),
    ],
}

BANG_MUC_7B = {
    "tieu_de": ("Bảng 7.2. Chi phí trên mờ chuyển động, cùng năm mức phần dư biến phân. "
                "Dòng cuối được bàn riêng trong phần chữ."),
    "cot": COT_MUC_7,
    "dong": [
        ("3,0·10⁻²", "5786", "3,97", "245", "0,55", "23,62", "7,15"),
        ("2,0·10⁻²", "5848", "6,57", "321", "0,76", "18,22", "8,60"),
        ("1,5·10⁻²", "8448", "10,98", "434", "0,91", "19,47", "12,08"),
        ("1,2·10⁻²", "8494", "15,33", "493", "1,04", "17,23", "14,72"),
        ("1,0·10⁻²", "435795", "732,89", "773", "1,54", "563,77", "476,69"),
    ],
}

MUC_7 = [
    ("Mục 7. Ngân sách thích nghi và kết quả chi phí", "de_muc"),

    ("Mục 6 kết thúc bằng một chứng chỉ mà chương trình kiểm được ở mỗi bước nội. Mục "
     "này dùng chứng chỉ đó để xây một chế độ chọn ngân sách bước nội, rồi đo xem chế "
     "độ ấy tiết kiệm được bao nhiêu so với phép chiếu chính xác. Đây là phần kết quả "
     "số chính của giai đoạn nghiên cứu.", "thuong"),

    ("7.1. Chế độ ngân sách thích nghi", "de_muc_phu"),
    ("Chế độ được mô tả như sau. Trước khi chạy, đặt một lịch sai số giảm dần theo số "
     "bước ngoài: mức sai số cho phép ở bước ngoài thứ k bằng hệ số đầu chia cho k cộng "
     "một, tất cả lũy thừa một số mũ lớn hơn một. Tại mỗi bước ngoài, vòng lặp nội chạy "
     "cho tới khi chứng chỉ xuống dưới mức của lịch tại bước ngoài đó, rồi dừng.",
     "thuong"),

    ("Có hai điểm cần giải thích trong cách đặt lịch. Thứ nhất, số mũ phải lớn hơn một "
     "vì chỉ khi đó chuỗi các mức sai số mới hội tụ, tức dãy sai số tổng được; đây đúng "
     "là giả thiết mà định lý ở mục 9 cần, nên lịch không phải chọn tùy tiện mà được "
     "chọn để thỏa giả thiết. Thứ hai, mức sai số giảm dần chứ không cố định, vì ở những "
     "bước ngoài đầu nghiệm còn xa nên chiếu chính xác là lãng phí, còn ở những bước "
     "ngoài sau thì cần chính xác hơn để không làm hỏng hội tụ.", "thuong"),

    ("Vòng lặp nội dùng khởi tạo ấm: biến gốc và biến đối ngẫu của bước ngoài trước "
     "được giữ lại làm điểm xuất phát cho bước ngoài sau. Nhờ vậy phần lớn bước ngoài "
     "chỉ tốn vài bước nội, vì điểm xuất phát đã gần nghiệm chiếu mới.", "thuong"),

    ("7.2. Thiết lập thực nghiệm", "de_muc_phu"),
    ("Bài toán thử nghiệm là khử mờ ảnh. Ảnh xám cạnh 96 điểm ảnh, tám ảnh kiểm tra, "
     "hai loại mờ là mờ Gauss và mờ chuyển động, nhiễu cộng có độ lệch chuẩn 0,05. Mỗi "
     "lần chạy gồm 150 bước ngoài. Bán kính quả cầu biến phân toàn phần đặt bằng 0,55 "
     "lần biến phân toàn phần của ảnh gốc, giống nhau cho mọi cấu hình.", "thuong"),

    ("Hằng số Lipschitz của toán tử được ước lượng bằng phép lặp lũy thừa, và độ dài "
     "bước lấy bằng 0,9 lần cận trên lý thuyết. Đại lượng dùng để so sánh là phần dư "
     "biến phân, tức chuẩn của phần dư bất đẳng thức biến phân sau khi chuẩn hóa. Toàn "
     "bộ phép đo chạy trên máy chủ có bộ xử lý đồ họa, với đồng bộ hóa trước và sau mỗi "
     "phép bấm giờ; lý do phải đồng bộ hóa được nêu ở mục 8.", "thuong"),

    ("7.3. Kết quả chi phí", "de_muc_phu"),
    ("Cách so sánh như sau. Năm mức phần dư biến phân được ấn định trước khi chạy. Với "
     "mỗi mức, ghi lại chi phí nhỏ nhất mà phép chiếu chính xác cần để đạt mức đó, và "
     "chi phí nhỏ nhất mà chế độ thích nghi cần để đạt cùng mức đó. Chi phí đo bằng hai "
     "thước đo song song: tổng số bước nội và thời gian chạy thuật toán. Kết quả ở bảng "
     "7.1 và bảng 7.2.", "thuong"),

    ("", "bang7a"),
    ("", "bang7b"),

    ("Trên mờ Gauss, chế độ thích nghi rẻ hơn phép chiếu chính xác từ 13,2 đến 19,2 lần "
     "theo bước nội, và từ 4,7 đến 9,7 lần theo thời gian. Trên mờ chuyển động, ở bốn "
     "mức phần dư biến phân đầu, các con số tương ứng là từ 17,2 đến 23,6 lần và từ 7,2 "
     "đến 14,7 lần.", "thuong"),

    ("Cần chú ý là hệ số theo thời gian luôn nhỏ hơn hệ số theo bước nội. Lý do là một "
     "bước nội của chế độ thích nghi đắt hơn một bước nội của phép chiếu chính xác, vì "
     "nó còn phải tính chứng chỉ. Đây chính là chỗ mà cách đo chỉ đếm bước nội sẽ thổi "
     "phồng lợi thế, và mục 8 sẽ bàn kỹ.", "thuong"),

    ("Dòng cuối của bảng 7.2 phải được đọc riêng và không nên trích như một thành tích. "
     "Ở mức phần dư biến phân chặt nhất trên mờ chuyển động, phép chiếu chính xác cần "
     "435795 bước nội và 732,89 giây, khiến hệ số vọt lên 563,77 lần theo bước nội và "
     "476,69 lần theo thời gian. Con số lớn ấy phản ánh việc phép chiếu chính xác gần "
     "như không đạt được mức đó trong ngân sách cho phép, chứ không phản ánh chế độ "
     "thích nghi mạnh lên. Khoảng nên trích dẫn vẫn là khoảng của bốn mức đầu.",
     "thuong"),

    ("7.4. Tính khả thi của đầu ra", "de_muc_phu"),
    ("Một kết quả đi kèm đáng nêu là mọi đầu ra đều khả thi. Trên toàn bộ 92 cấu hình "
     "của hai lưới thực nghiệm, biến phân toàn phần của ảnh đầu ra chia cho bán kính "
     "quả cầu không bao giờ vượt 1,0000, tức không cấu hình nào ra ngoài tập ràng "
     "buộc; hai cấu hình còn dừng hẳn bên trong, ở mức 0,9994 và 0,9995. Điều này có "
     "được là nhờ phép ép về tập ràng buộc mô tả ở mục 6.3, vốn được gắn liền với việc "
     "tính chứng chỉ.", "thuong"),

    ("Để thấy điều này không hiển nhiên, có thể so với chế độ ngân sách cố định. Khi "
     "mỗi bước ngoài chỉ chạy đúng một bước nội, đầu ra trên mờ chuyển động có biến "
     "phân toàn phần bằng 1,2433 lần bán kính, tức vượt ra ngoài tập ràng buộc 24,3 "
     "phần trăm. Một nghiệm như vậy không còn thỏa bài toán ban đầu, dù các chỉ số chất "
     "lượng ảnh của nó trông không tệ.", "thuong"),

    ("7.5. Chỗ phương pháp thua", "de_muc_phu"),
    ("Chế độ thích nghi có một điểm yếu rõ ràng: nó rất nhạy với lịch sai số, và một "
     "lịch đặt sai làm chi phí bùng nổ. Phần này trình bày điểm yếu đó ngang hàng với "
     "phần thắng ở trên.", "thuong"),

    ("Trên lưới lịch đầy đủ gồm mười sáu cấu hình cho mỗi loại mờ, cấu hình xấu nhất ở "
     "cả hai loại mờ đều là cấu hình có hệ số đầu nhỏ nhất và số mũ lớn nhất. Nó tốn "
     "hơn mốc phép chiếu chính xác 42,83 lần trên mờ Gauss và 39,64 lần trên mờ chuyển "
     "động. Ngay trong lưới công bằng đã lọc bớt lịch xấu, khoảng cách giữa cấu hình "
     "đắt nhất và cấu hình rẻ nhất vẫn là 970 lần trên mờ Gauss và 1104 lần trên mờ "
     "chuyển động.", "thuong"),

    ("Nguyên nhân là như sau. Số mũ lớn làm mức sai số cho phép giảm quá nhanh, nên chỉ "
     "sau ít bước ngoài lịch đã đòi độ chính xác gần bằng phép chiếu chính xác, mà khi "
     "ấy chi phí mỗi bước ngoài tăng vọt trong khi chất lượng nghiệm gần như không đổi. "
     "Hệ số đầu nhỏ gây ra cùng một hiệu ứng ngay từ bước ngoài đầu tiên.", "thuong"),

    ("Quy luật chọn lịch rút ra từ đây rất đơn giản: lấy số mũ sát một từ phía trên, và "
     "lấy hệ số đầu lớn. Trong các thực nghiệm trên, số mũ 1,01 cùng hệ số đầu từ 4 đến "
     "8 cho kết quả tốt nhất ở cả hai loại mờ. Đây là một khuyến nghị thực hành rút từ "
     "số liệu, không phải một kết quả lý thuyết.", "thuong"),

    ("7.6. Dãy sai số có tổng được hay không", "de_muc_phu"),
    ("Mục 5.3 đã hoãn lại lý do thứ hai để bỏ bước neo; phần này trả lời. Câu hỏi là "
     "dãy sai số của phép chiếu xấp xỉ dọc quỹ đạo sinh ra có tổng được hay không, vì "
     "đó là giả thiết mà định lý ở mục 9 cần.", "thuong"),

    ("Phép đo thực hiện trên 3000 bước ngoài, dài gấp hai mươi lần các thực nghiệm ở "
     "trên, và đo độ dốc của dịch chuyển bước chiếu trên thang lôgarit hai chiều. Khi "
     "không dùng bước neo, độ dốc là âm 2,342 trên mờ Gauss và âm 2,838 trên mờ chuyển "
     "động. Cả hai đều nhỏ hơn âm một khá xa, nên chuỗi hội tụ; tổng tích lũy đến hết "
     "quỹ đạo chỉ lớn hơn tổng tích lũy đến giữa quỹ đạo 1,006 lần và 1,000 lần, tức "
     "đã bão hòa.", "thuong"),

    ("Khi thêm bước neo, độ dốc tụt xuống âm 1,191 trên mờ Gauss và âm 1,011 trên mờ "
     "chuyển động. Con số thứ hai đáng lo: ngưỡng để chuỗi còn hội tụ là âm một, nên "
     "âm 1,011 nằm sát ngay ranh giới. Với một biên hẹp như vậy, không thể khẳng định "
     "chắc chắn giả thiết tổng được vẫn đúng khi đổi bài toán hoặc đổi tham số.",
     "thuong"),

    ("Đây là lý do thứ hai để bỏ bước neo, bên cạnh lý do về chuỗi phép chiếu đã nêu ở "
     "mục 5.3. Sau khi bỏ bước neo, giả thiết tổng được không những đúng mà còn đúng "
     "với biên rộng, và điều đó làm phần lý thuyết ở mục 9 đứng vững hơn hẳn.",
     "thuong"),

    ("Tất cả các con số trong mục này chỉ có nghĩa nếu phép đo công bằng. Mục 8 trình "
     "bày giao thức đo đã dùng, và ba lỗi đo lường đã phải tự sửa trước khi có được "
     "các con số trên.", "thuong"),
]


# Bảng ở mục 8: bẫy đo lường khi áp lịch bước tăng tốc, trích từ compare_accel_gauss.log.
BANG_MUC_8 = {
    "tieu_de": ("Bảng 8.1. Bẫy đo lường: áp lịch bước tăng tốc cho mọi chế độ làm hệ số "
                "tiết kiệm trông lớn hơn, nhưng là do phương pháp đối chứng chậm đi."),
    "cot": ["Chế độ", "Bước nội khi không tăng tốc", "Bước nội khi tăng tốc",
            "Hệ số tiết kiệm khi không tăng tốc", "Hệ số tiết kiệm khi tăng tốc"],
    "dong": [
        ("Phép chiếu chính xác", "384", "560", "mốc so sánh", "mốc so sánh"),
        ("Ngân sách cố định hai bước nội", "162", "162", "2,37", "3,46"),
        ("Ngân sách cố định năm bước nội", "255", "305", "1,51", "1,84"),
        ("Ngân sách theo lôgarit", "289", "379", "1,33", "1,48"),
    ],
}

MUC_8 = [
    ("Mục 8. Giao thức đo chi phí công bằng", "de_muc"),

    ("Đóng góp của mục 7 chủ yếu là một so sánh chi phí, mà một so sánh chi phí chỉ đáng "
     "tin khi phép đo công bằng. Mục này trình bày giao thức đo đã dùng. Cách trình bày "
     "là kể lại ba lỗi đo lường đã mắc rồi tự phát hiện và sửa, cùng một lỗi kỹ thuật và "
     "một cái bẫy đã kịp tránh. Viết như vậy vì các con số trước khi sửa đã từng được "
     "ghi vào tài liệu nội bộ, nên việc giấu chúng đi sẽ làm sai lệch bức tranh về quá "
     "trình nghiên cứu.", "thuong"),

    ("8.1. Lỗi thứ nhất: chỉ đếm bước nội mà không đo thời gian", "de_muc_phu"),
    ("Bản đo đầu tiên chỉ đếm tổng số bước nội. Theo thước đo đó, chế độ thích nghi rẻ "
     "hơn phép chiếu chính xác 12,86 lần trên mờ Gauss và 14,90 lần trên mờ chuyển "
     "động. Hai con số này đã được ghi lại và đã suýt được dùng làm kết quả chính.",
     "thuong"),

    ("Chúng thổi phồng lợi thế. Một bước nội của chế độ thích nghi không tương đương một "
     "bước nội của phép chiếu chính xác, vì nó còn phải tính chứng chỉ, tức tính thêm "
     "hai hàm mục tiêu ở mỗi bước. Đếm hai loại bước nội như nhau là so sánh hai đơn vị "
     "khác nhau.", "thuong"),

    ("Cách sửa là báo song song cả hai thước đo, số bước nội và thời gian chạy thuật "
     "toán, như hai cột cuối của bảng 7.1 và bảng 7.2. Sau khi sửa, hệ số theo thời gian "
     "thấp hơn hệ số theo bước nội một cách nhất quán, đúng như dự đoán. Con số nên "
     "trích dẫn là con số theo thời gian, vì đó là thứ người dùng thật sự phải trả.",
     "thuong"),

    ("8.2. Lỗi thứ hai: mức phần dư biến phân mục tiêu chọn sau khi đã thấy số liệu",
     "de_muc_phu"),
    ("Bản đo đầu tiên lấy mức phần dư biến phân mục tiêu từ chính kết quả cuối của "
     "phương pháp đối chứng, rồi hỏi chế độ thích nghi cần bao nhiêu chi phí để đạt "
     "mức đó. Cách làm này có vẻ tự nhiên nhưng là một vòng luẩn quẩn: mục tiêu được "
     "đặt sau khi đã nhìn thấy số liệu, nên có thể vô tình rơi vào đúng vùng mà phương "
     "pháp đề xuất có lợi.", "thuong"),

    ("Cách sửa là ấn định trước năm mức phần dư biến phân, ghi thẳng vào mã nguồn trước "
     "khi chạy, và giữ nguyên cho cả hai loại mờ lẫn mọi cấu hình. Năm mức ấy là các "
     "mức trong cột đầu của bảng 7.1 và bảng 7.2. Cấu hình nào không đạt mức nào thì bị "
     "loại khỏi mức đó, không có ngoại lệ.", "thuong"),

    ("8.3. Lỗi thứ ba: dò tham số bất đối xứng", "de_muc_phu"),
    ("Bản đo đầu tiên dò mười sáu cấu hình cho phương pháp đề xuất nhưng chỉ hai cấu "
     "hình cho phương pháp đối chứng. Khi một bên được dò kỹ hơn bên kia tám lần thì "
     "phần thắng có thể chỉ là phần thưởng của công dò, chứ không phải ưu thế của "
     "phương pháp.", "thuong"),

    ("Cách sửa là dò bằng nhau ở hai nhóm chính: tám cấu hình cho nhóm thích nghi và "
     "tám cấu hình cho nhóm phép chiếu chính xác. Ngoài ra có thêm hai nhóm tiêu chuẩn "
     "tương đối làm đối chứng phụ, với bốn và tám cấu hình. Với mỗi nhóm, chỉ cấu hình "
     "tốt nhất tại mỗi mức phần dư biến phân được đưa vào bảng, nên mỗi nhóm đều được "
     "trình bày ở trạng thái mạnh nhất của nó.", "thuong"),

    ("8.4. Một lỗi kỹ thuật khi bấm giờ", "de_muc_phu"),
    ("Ngoài ba lỗi trên còn một lỗi kỹ thuật. Bộ xử lý đồ họa chạy bất đồng bộ với bộ "
     "xử lý trung tâm: lệnh được xếp vào hàng đợi rồi trả về ngay, trước khi phép tính "
     "thật sự xong. Bấm giờ mà không đồng bộ hóa thì chỉ đo thời gian xếp lệnh chứ "
     "không đo thời gian tính, và con số thu được gần như vô nghĩa.", "thuong"),

    ("Cách sửa là gọi đồng bộ hóa ngay trước và ngay sau mỗi đoạn được bấm giờ. Cùng "
     "với đó, thời gian thuật toán được tách khỏi thời gian đo đạc phụ trợ như tính chỉ "
     "số chất lượng ảnh hay ghi tệp. Việc tách này quan trọng hơn vẻ ngoài của nó: phần "
     "đo đạc phụ trợ tốn khoảng 45 giây mỗi lần chạy, đủ để nhấn chìm phần thuật toán "
     "vốn chỉ khoảng một giây, và nếu lấy cột thời gian tổng thì mọi khác biệt giữa các "
     "chế độ sẽ biến mất.", "thuong"),

    ("8.5. Một cái bẫy đã tránh", "de_muc_phu"),
    ("Có một thay đổi từng làm kết quả trông đẹp hơn mà cuối cùng bị loại. Bộ giải nội "
     "Chambolle-Pock có một lịch bước tăng tốc dành cho bài toán lồi mạnh. Khi áp lịch "
     "này cho mọi chế độ, hệ số tiết kiệm của chế độ ngân sách cố định hai bước nội "
     "tăng từ 2,37 lên 3,46 lần.", "thuong"),

    ("", "bang8"),

    ("Nhìn vào bảng 8.1 sẽ thấy vì sao con số đó không dùng được. Chế độ ngân sách cố "
     "định hai bước nội vẫn tốn đúng 162 bước nội như cũ, tức nó không hề tốt lên. Cái "
     "thay đổi là phép chiếu chính xác, vốn dùng làm mốc so sánh, tăng từ 384 lên 560 "
     "bước nội. Hệ số tăng lên hoàn toàn vì mẫu số xấu đi.", "thuong"),

    ("Lịch bước tăng tốc do đó bị loại, và kết quả âm tính này được ghi lại ở mục 10. "
     "Bài học rút ra là khi một thay đổi làm hệ số so sánh tăng lên, phải kiểm cả tử số "
     "lẫn mẫu số trước khi mừng.", "thuong"),

    ("8.6. Kết luận của mục", "de_muc_phu"),
    ("Con số đáng tin là con số đo sau khi đã sửa cả ba lỗi và lỗi kỹ thuật nói trên, "
     "tức các con số trong bảng 7.1 và bảng 7.2. Chúng nhỏ hơn các con số ban đầu, "
     "nhưng chúng chịu được phản biện. Giao thức này cũng là câu trả lời cho nguyên tắc "
     "thứ nhất và nguyên tắc thứ hai ở mục 3, vốn đòi đặt tiêu chí trước khi chạy và "
     "cho phương pháp đối chứng cơ hội mạnh nhất.", "thuong"),

    ("Phần thực nghiệm đến đây là hết. Mục 9 chuyển sang phần lý thuyết, và trình bày "
     "định lý hội tụ cùng những giả thiết mà mục 7 đã kiểm bằng số.", "thuong"),
]


MUC_9 = [
    ("Mục 9. Phần lý thuyết", "de_muc"),

    ("Mục 7 và mục 8 trình bày phần thực nghiệm. Mục này chuyển sang phần lý thuyết và "
     "trả lời hai câu hỏi: đã chứng minh được gì, và chưa chứng minh được gì. Cách "
     "trình bày tách bạch hai phần đó, vì trong quá trình làm việc đã có hai bản thảo "
     "chứng minh bị bác, và bài học rút ra là không nên gộp phần đã vững với phần còn "
     "phác thảo.", "thuong"),

    ("9.1. Bài toán và sơ đồ cuối cùng", "de_muc_phu"),
    ("Bài toán là bất đẳng thức biến phân trên tập ràng buộc: tìm một điểm của tập ràng "
     "buộc sao cho tích vô hướng giữa toán tử chi phí tại điểm đó với hướng đi tới mọi "
     "điểm khác của tập ràng buộc đều không âm. Tập các điểm như vậy gọi là tập nghiệm "
     "và được giả thiết khác rỗng.", "thuong"),

    ("Sơ đồ sau khi đã bỏ bước quán tính và bước neo ở mục 5 chỉ còn hai dòng. Dòng thứ "
     "nhất lấy điểm phản xạ bằng hai lần điểm lặp hiện tại trừ điểm lặp trước đó. Dòng "
     "thứ hai đi một bước theo hướng ngược với toán tử chi phí tại điểm phản xạ, với độ "
     "dài bước cố định, rồi chiếu xấp xỉ điểm thu được lên tập ràng buộc với mức sai số "
     "cho phép của bước ngoài đó. Toán tử chi phí chỉ được tính một lần trong mỗi bước "
     "ngoài.", "thuong"),

    ("9.2. Các giả thiết", "de_muc_phu"),
    ("Định lý dùng năm giả thiết. Thứ nhất, tập ràng buộc lồi, đóng và khác rỗng, và "
     "tập nghiệm khác rỗng. Thứ hai, toán tử chi phí đơn điệu và liên tục Lipschitz. "
     "Thứ ba, độ dài bước dương và nhỏ hơn một thương, mà tử số là căn bậc hai của hai "
     "trừ đi một, còn mẫu số là hằng số Lipschitz. Thứ tư, phép chiếu xấp xỉ luôn trả "
     "về điểm nằm trong tập ràng buộc. Thứ năm, dãy sai số của phép chiếu xấp xỉ tổng "
     "được.", "thuong"),

    ("Giả thiết thứ hai cần một lưu ý mà nếu bỏ qua thì chứng minh sai. Tính đơn điệu "
     "và tính Lipschitz phải được đặt trên toàn không gian, chứ không chỉ trên tập ràng "
     "buộc. Lý do nằm ở chính điểm phản xạ: vì nó bằng hai lần điểm lặp hiện tại trừ "
     "điểm lặp trước, nó nói chung nằm ngoài tập ràng buộc, dù cả hai điểm lặp đều nằm "
     "trong. Trong chứng minh, tính đơn điệu được dùng đúng một lần, và dùng tại cặp "
     "gồm điểm phản xạ với một điểm nghiệm; nếu chỉ giả thiết đơn điệu trên tập ràng "
     "buộc thì bước đó không hợp lệ.", "thuong"),

    ("Với bài toán khôi phục ảnh của đề tài, giả thiết này thỏa: toán tử chi phí có "
     "dạng chuyển vị của toán tử làm mờ nhân với phần dư, nên nó đơn điệu và liên tục "
     "Lipschitz trên toàn không gian.", "thuong"),

    ("9.3. Vai trò của giả thiết khả thi", "de_muc_phu"),
    ("Giả thiết thứ tư, rằng phép chiếu xấp xỉ trả về điểm khả thi, dễ bị hiểu nhầm là "
     "một chi tiết cài đặt. Thực ra nó là điều kiện của định lý, và nếu bỏ đi thì chứng "
     "minh hỏng ở một chỗ cụ thể.", "thuong"),

    ("Chỗ đó như sau. Trong chứng minh có một đại lượng bằng tích vô hướng giữa toán tử "
     "chi phí tại điểm nghiệm với hướng đi từ điểm lặp tới điểm nghiệm. Đại lượng này "
     "không dương đúng khi điểm lặp nằm trong tập ràng buộc, vì đó chính là phát biểu "
     "của bất đẳng thức biến phân tại điểm nghiệm. Dấu không dương ấy được dùng hai "
     "lần: một lần để khẳng định đại lượng thế năng không âm, và một lần để giữ dấu của "
     "số hạng thu gọn dần khi cộng dồn qua các bước ngoài. Nếu điểm lặp ra ngoài tập "
     "ràng buộc thì dấu này mất, và cả hai lập luận đổ theo.", "thuong"),

    ("Đây cũng là chỗ phần thực nghiệm gặp phần lý thuyết. Con số ở mục 7.4, rằng biến "
     "phân toàn phần của đầu ra chia cho bán kính không bao giờ vượt 1,0000 trên toàn "
     "bộ cấu hình đã chạy, không phải một chỉ số phụ mà chính là giả thiết thứ tư được "
     "kiểm bằng số. Tương tự, phép đo tổng được ở mục 7.6 chính là giả thiết thứ năm "
     "được kiểm bằng số. Hai giả thiết của định lý vì thế không phải điều kiện đặt ra "
     "cho tiện, mà là điều kiện đã được đo.", "thuong"),

    ("9.4. Phần đã chứng minh đầy đủ", "de_muc_phu"),
    ("Phần này gồm bốn bổ đề và một bước ghép, tất cả đã được ba người phản biện độc "
     "lập tính lại từng phép biến đổi và đều xác nhận đúng, không sai dấu, sai hệ số "
     "hay sai chỉ số.", "thuong"),

    ("Bổ đề thứ nhất là bất đẳng thức cơ bản của bước chiếu, dẫn từ đặc trưng của phép "
     "chiếu lên tập lồi đóng cộng với đẳng thức phân cực. Bổ đề thứ hai xử lý số hạng "
     "toán tử bằng tính đơn điệu; đây là chỗ duy nhất dùng tính đơn điệu, và là lý do "
     "giả thiết thứ hai phải đặt trên toàn không gian. Bổ đề thứ ba thu gọn số hạng "
     "nghiệm: nhờ dạng của điểm phản xạ, đại lượng tại điểm phản xạ bằng hai lần đại "
     "lượng tại bước hiện tại trừ đại lượng tại bước trước, nên khi cộng dồn qua các "
     "bước ngoài phần lớn triệt tiêu. Bổ đề thứ tư chặn ảnh hưởng của sai số phép chiếu "
     "xấp xỉ bằng bất đẳng thức tam giác và khai triển bình phương.", "thuong"),

    ("Bước ghép dựng một đại lượng thế năng gồm ba số hạng: bình phương khoảng cách từ "
     "điểm lặp tới điểm nghiệm, một số hạng bình phương đo độ lệch giữa điểm lặp với "
     "điểm phản xạ của bước trước, và số hạng nghiệm của bước trước lấy dấu ngược. Cả "
     "ba đều không âm, số hạng thứ ba nhờ đúng giả thiết khả thi nói ở trên.", "thuong"),

    ("Trong bước ghép có một đẳng thức đáng nêu vì nó là chỗ mọi thứ khớp lại. Hai hằng "
     "số xuất hiện trong bất đẳng thức một bước, một hằng số nhân với bình phương độ "
     "lệch giữa hai điểm lặp liên tiếp và một hằng số nhân với bình phương độ lệch giữa "
     "điểm chiếu với điểm phản xạ, thoạt nhìn khác nhau. Nhưng hằng số thứ hai trừ đi "
     "tích của độ dài bước với hằng số Lipschitz thì đúng bằng hằng số thứ nhất. Nhờ "
     "đẳng thức đó, hai số hạng âm gộp lại với cùng một hệ số, và đại lượng thế năng "
     "giảm đơn điệu. Điều kiện để cả hai hằng số dương chính là giả thiết thứ ba về độ "
     "dài bước.", "thuong"),

    ("Với phép chiếu chính xác, lập luận đến đây là khép kín: đại lượng thế năng giảm "
     "và bị chặn dưới bởi không nên hội tụ, và tổng các số hạng âm hữu hạn.", "thuong"),

    ("9.5. Nhiễu do phép chiếu xấp xỉ", "de_muc_phu"),
    ("Khi phép chiếu chỉ xấp xỉ, nhiễu vào theo hai đường khác nhau, và việc phân biệt "
     "chúng là phần khó nhất của chứng minh.", "thuong"),

    ("Đường thứ nhất là nhiễu điểm cuối. Điểm lặp thực khác điểm chiếu chính xác lý "
     "tưởng một lượng không quá mức sai số cho phép, nên khi thay điểm này bằng điểm "
     "kia trong đại lượng thế năng sẽ sinh các số hạng bậc nhất và bậc hai theo mức sai "
     "số. Đường này thuần túy kỹ thuật, và phản biện xác nhận đúng tới từng hệ số.",
     "thuong"),

    ("Đường thứ hai là nhiễu điểm gốc, và tinh tế hơn. Bất đẳng thức một bước có một số "
     "hạng còn phụ thuộc điểm lặp cách đó hai bước, và số hạng ấy được chặn bằng đặc "
     "trưng phép chiếu của bước trước. Nhưng đặc trưng đó chỉ đúng cho điểm chiếu chính "
     "xác của bước trước, chứ không cho điểm lặp thực. Chênh lệch giữa hai điểm ấy sinh "
     "thêm các số hạng chéo bậc nhất theo sai số của bước trước.", "thuong"),

    ("Câu hỏi quyết định là hằng số đứng trước sai số của bước trước lớn cỡ nào. Nếu nó "
     "chứa nghịch đảo của độ dài bước thì sai số bị khuếch đại khi độ dài bước nhỏ, và "
     "đó sẽ là một hiện tượng riêng của sơ đồ này, tức một điểm phân biệt. Hằng số này "
     "đã được dẫn xuất đầy đủ, và kết quả là nó gồm bốn phần: hai lần chặn của độ lệch "
     "giữa hai điểm lặp liên tiếp, hai lần tích của độ dài bước với chặn của toán tử "
     "chi phí, hai lần chặn của khoảng cách từ điểm lặp tới điểm so sánh, và hai lần "
     "mức sai số ban đầu.", "thuong"),

    ("9.6. Kết luận trung thực của phần lý thuyết", "de_muc_phu"),
    ("Hằng số vừa nêu không chứa nghịch đảo của độ dài bước. Ngược lại, độ dài bước xuất "
     "hiện ở một phần dưới dạng thừa số nhân, tức làm phần đó nhỏ đi khi độ dài bước "
     "nhỏ. Lý do có tính cấu trúc: trong đặc trưng phép chiếu, toán tử chi phí đã đi "
     "kèm sẵn thừa số độ dài bước, nên khi sai số điểm gốc gặp số hạng toán tử, nó nhận "
     "thừa số ấy chứ không nhận nghịch đảo của nó.", "thuong"),

    ("Khẳng định ban đầu của báo cáo, rằng sai số điểm gốc bị khuếch đại bởi nghịch đảo "
     "độ dài bước, vì thế là sai; phản biện đã nghi ngờ đúng. Hệ quả phải nói "
     "thẳng: số hạng nhiễu điểm gốc không có gì đặc biệt hơn nhiễu điểm cuối, cả hai "
     "đều bậc nhất theo sai số và đều tổng được. Định lý hội tụ của báo cáo do đó là "
     "một mở rộng của kết quả đã có, chứ không phải một cơ chế mới. Việc thêm bước phản "
     "xạ vào khung phép chiếu xấp xỉ chỉ là kiểm rằng cơ chế bền vững với nhiễu vẫn "
     "chạy qua bước phản xạ, và kết quả kiểm là đúng.", "thuong"),

    ("Đây chính là lý do bài báo nhắm nhóm Q2 chứ không phải nhóm Q1. Kết luận này "
     "không làm mất giá trị của công việc mà xác định đúng giá trị của nó: đóng góp nằm "
     "ở chứng chỉ sai số tính được ở mục 6, chế độ ngân sách thích nghi và phân tích "
     "chi phí ở mục 7, cùng giao thức đo ở mục 8. Ba thứ đó là đóng góp về thuật toán "
     "và thực nghiệm, và chúng không phụ thuộc vào việc phần định lý có mới hay không.",
     "thuong"),

    ("9.7. Hai chỗ còn phải viết đầy đủ", "de_muc_phu"),
    ("Chỗ thứ nhất là phát biểu chính xác của bổ đề tựa Fejér. Sau khi thêm nhiễu, "
     "truy hồi không còn dạng đại lượng thế năng của bước sau không vượt đại lượng thế "
     "năng của bước trước cộng một sai số cộng tính, mà có thêm một thừa số nhân lớn "
     "hơn một. Dạng nhân tính này vẫn cho hội tụ khi cả dãy hệ số nhân lẫn dãy sai số "
     "cộng đều tổng được, nhưng phải dùng đúng phát biểu dành cho dạng đó, và phải kiểm "
     "rằng bước biến đổi không làm mất phần âm vốn là thứ cho ra tính tổng được của các "
     "độ lệch.", "thuong"),

    ("Chỗ thứ hai là bước chuyển qua giới hạn yếu. Lập luận cuối lấy một dãy con hội tụ "
     "yếu, chuyển qua giới hạn trong đặc trưng phép chiếu, rồi dùng bổ đề Minty để kết "
     "luận điểm giới hạn thuộc tập nghiệm, sau đó dùng bổ đề Opial. Trong không gian "
     "hữu hạn chiều của phần thực nghiệm, các bước này là thường quy. Trong không gian "
     "vô hạn chiều, bước Minty còn cần tính liên tục yếu theo dãy của toán tử chi phí, "
     "và điều đó chưa được kiểm.", "thuong"),

    ("Tóm lại, phần lý thuyết đứng vững ở khối chính và còn hai chỗ kỹ thuật phải viết "
     "cho đủ. Mục 10 tiếp theo ghi lại các hướng đã thử mà không đi tới đâu, trong đó "
     "có cả hai bản thảo chứng minh đã bị bác nói ở đầu mục này.", "thuong"),
]


MUC_10 = [
    ("Mục 10. Các kết quả âm tính đã ghi nhận", "de_muc"),

    ("Mục này ghi lại bốn hướng đã thử và đã đóng lại. Lý do dành hẳn một mục cho chúng "
     "là để người đọc, kể cả chính đề tài về sau, không mất công thử lại. Bốn hướng "
     "được trình bày theo cùng một khuôn: ý tưởng là gì, vì sao lúc đầu nó có vẻ hợp "
     "lý, số liệu đo được ra sao, và nguyên nhân thất bại nằm ở đâu. Hai trong bốn "
     "hướng đã được nhắc tới ở mục 6 và mục 8; ở đây chúng được ghi đầy đủ.", "thuong"),

    ("10.1. Lịch bước tăng tốc cho bài toán chiếu", "de_muc_phu"),
    ("Ý tưởng: thuật toán Chambolle-Pock có một lịch bước tăng tốc dành riêng cho bài "
     "toán có hàm mục tiêu lồi mạnh, trong đó độ dài bước gốc co dần còn độ dài bước "
     "đối ngẫu giãn dần. Bài toán chiếu lên quả cầu biến phân toàn phần đúng là lồi "
     "mạnh, nên lịch này áp được và trên lý thuyết cho tốc độ hội tụ tốt hơn.",
     "thuong"),

    ("Số liệu: không chế độ nào nhanh lên. Một chế độ giữ nguyên số bước nội, ba chế độ "
     "còn lại tốn thêm bước nội. Nếu lấy tỉ số giữa chi phí khi không tăng tốc và chi "
     "phí khi tăng tốc, thì tỉ số đó là 0,84, 0,76 và 0,69, tức đều nhỏ hơn một nên đều "
     "là chậm đi; chế độ chậm đi nhiều nhất chính là phép chiếu chính xác.", "thuong"),

    ("Nguyên nhân: lịch tăng tốc co độ dài bước gốc ngay từ những bước nội đầu tiên, "
     "trong khi chế độ của báo cáo chỉ chạy vài bước nội mỗi bước ngoài nhờ khởi tạo "
     "ấm. Lợi ích tiệm cận của lịch tăng tốc chỉ xuất hiện sau nhiều bước nội, còn cái "
     "giá là bước ngắn thì phải trả ngay. Thêm nữa, lịch được đặt lại ở mỗi bước ngoài, "
     "nên nó không bao giờ chạy đủ dài để lợi ích kịp xuất hiện. Kết quả này còn dẫn "
     "tới một bẫy đo lường đã trình bày ở mục 8.5.", "thuong"),

    ("10.2. Siết chứng chỉ bằng cách theo dõi giá trị tốt nhất", "de_muc_phu"),
    ("Ý tưởng: chứng chỉ ở mục 6 được tính từ giá trị hai hàm mục tiêu tại bước nội "
     "hiện tại. Nhưng chặn trên vẫn hợp lệ nếu thay bằng giá trị tốt nhất đã gặp trong "
     "các bước nội đã qua, tức giá trị nhỏ nhất của hàm mục tiêu gốc và giá trị lớn "
     "nhất của hàm mục tiêu đối ngẫu. Vì khoảng cách đối ngẫu tính theo cặp giá trị tốt "
     "nhất không lớn hơn khoảng cách tính theo cặp giá trị hiện tại, chứng chỉ chỉ có "
     "thể chặt hơn, và vòng lặp nội sẽ dừng sớm hơn.", "thuong"),

    ("Số liệu: không giảm được bước nội nào, ở mọi mức sai số đã thử.", "thuong"),

    ("Nguyên nhân: lập luận đúng nhưng vô ích trên thực tế, vì thuật toán Chambolle-Pock "
     "vốn đã cho hai dãy giá trị gần đơn điệu. Khi giá trị hiện tại gần như luôn là giá "
     "trị tốt nhất, việc theo dõi thêm giá trị tốt nhất chỉ tốn bộ nhớ và vài phép so "
     "sánh mà không đổi thời điểm dừng. Đây là một cải tiến đúng về lý thuyết nhưng "
     "không có chỗ để phát huy.", "thuong"),

    ("10.3. Tiêu chuẩn sai số tương đối theo chuẩn của toán tử", "de_muc_phu"),
    ("Ý tưởng: thay vì đặt trước một lịch sai số như ở mục 7.1, cho mức sai số cho phép "
     "tỉ lệ với chuẩn của toán tử chi phí tại điểm phản xạ. Cách này hấp dẫn vì nó tự "
     "điều chỉnh: khi còn xa nghiệm thì chuẩn lớn nên sai số cho phép lớn, khi gần "
     "nghiệm thì chuẩn nhỏ nên sai số cho phép tự siết lại, và người dùng không phải "
     "chọn lịch.", "thuong"),

    ("Nguyên nhân thất bại nằm ở lý thuyết, nên hướng này bị đóng trước khi đo. Định lý "
     "đòi mức sai số cho phép tiến về không, nhưng chuẩn của toán tử chi phí chỉ tiến "
     "về không khi nghiệm nằm hẳn bên trong tập ràng buộc. Nếu nghiệm nằm trên biên, "
     "tức có ràng buộc kích hoạt, thì tại nghiệm toán tử chi phí khác không. Với bài "
     "toán khôi phục ảnh dùng quả cầu biến phân toàn phần, nghiệm hầu như luôn nằm trên "
     "biên, vì bán kính được đặt nhỏ hơn biến phân toàn phần của ảnh gốc. Tiêu chuẩn "
     "này do đó hỏng đúng ở trường hợp mà đề tài quan tâm.", "thuong"),

    ("10.4. Tiêu chuẩn chiếu nới lỏng", "de_muc_phu"),
    ("Ý tưởng: đây là hướng được đầu tư nhiều nhất trong bốn hướng, và do một vòng phản "
     "biện đề xuất. Thay mô hình sai số dạng quả cầu bằng bất đẳng thức chiếu nới lỏng "
     "theo cách Díaz Millán định nghĩa, trong đó điểm trả về phải nằm trong tập ràng "
     "buộc và thỏa bất đẳng thức đặc trưng của phép chiếu với một sai số bậc hai theo "
     "dịch chuyển. Cái lợi rất lớn: sai số vào chứng minh ở bậc hai nên bị phần âm của "
     "bổ đề một bước hấp thụ, và khi đó dung sai chỉ cần là một hằng số chứ không cần "
     "một dãy tiến về không. Nếu làm được thì giả thiết thứ năm ở mục 9.2 biến mất.",
     "thuong"),

    ("Có đúng hai cách kiểm tiêu chuẩn này, và cả hai đã được thử. Cách thứ nhất là "
     "lấy cận trên của vế trái trên toàn tập ràng buộc. Ở đây gặp một rào cản: quả cầu "
     "biến phân toàn phần không bị chặn, vì mọi ảnh hằng đều có biến phân toàn phần "
     "bằng không; một ảnh hằng giá trị 1000 có biến phân toàn phần chỉ 2,56 nhân mười "
     "mũ trừ bốn trong khi chuẩn của nó là 16000. Rào cản này được gỡ bằng cách lấy tập "
     "ràng buộc là quả cầu biến phân toàn phần giao với hộp giá trị điểm ảnh, vốn là "
     "ràng buộc vật lý tự nhiên của ảnh; khi đó tập ràng buộc bị chặn.", "thuong"),

    ("Nhưng chỗ thất bại nằm ngay sau đó. Cận trên tính được rẻ duy nhất là thay tập "
     "ràng buộc bằng hộp bao ngoài, và cận này quá thô đến mức vô dụng. Đo trên bài "
     "chiếu thật, ngay cả khi chạy 2000 bước nội tức gần như chiếu chính xác, vế trái "
     "vẫn là 120,25 trong khi vế phải với dung sai 0,25 chỉ là 0,31. Tiêu chuẩn không "
     "bao giờ thỏa. Nguyên nhân có tính cấu trúc: quả cầu biến phân toàn phần giao hộp "
     "nhỏ hơn hộp rất nhiều, nên thay tập ràng buộc bằng hộp làm cận trên phồng lên; "
     "phép chiếu chính xác thỏa bất đẳng thức trên tập ràng buộc thật nhưng không thỏa "
     "trên hộp. Muốn kiểm đúng phải tính hàm tựa của tập ràng buộc, tức giải thêm một "
     "bài toán tối ưu lồi ở mỗi lần kiểm, đắt ngang chính phép chiếu.", "thuong"),

    ("Cách thứ hai là xuất bất đẳng thức nới lỏng từ chứng chỉ ở mục 6, thay vì kiểm "
     "bằng cận trên. Cách này chết vì sai bậc, và điều đó chứng minh được trên giấy: "
     "khai triển cho thấy vế trái bị chặn bởi mức sai số nhân một hằng số, tức bậc nhất "
     "theo mức sai số, trong khi tiêu chuẩn nới lỏng đòi vế trái không vượt một hằng số "
     "nhân bình phương dịch chuyển, tức bậc hai. Muốn ghép hai bên phải chọn mức sai số "
     "cỡ bình phương của dịch chuyển.", "thuong"),

    ("Số liệu xác nhận mức độ chênh. Tại bước ngoài thứ 199 của bài chiếu thật, dịch "
     "chuyển là 2,0236 nhân mười mũ trừ ba, trong khi mức sai số cần thiết với dung sai "
     "0,25 chỉ là 2,0894 nhân mười mũ trừ tám, tức phải siết chặt hơn dịch chuyển tới "
     "năm bậc. Tỉ lệ giữa hai đại lượng còn xấu đi khi tiến tới nghiệm, từ 8,8 nhân "
     "mười mũ trừ ba ở bước ngoài đầu xuống 1,0 nhân mười mũ trừ năm ở bước ngoài thứ "
     "199. Chi phí vòng lặp nội vì thế bùng nổ.", "thuong"),

    ("Kết luận cho hướng này: tiêu chuẩn chiếu nới lỏng đúng về lý thuyết nhưng không "
     "thực thi được trên bài toán chiếu lên quả cầu biến phân toàn phần với bộ giải nội "
     "hiện có. Chính vì cả hai cách đều đóng mà báo cáo quay lại mô hình sai số dạng "
     "quả cầu cùng giả thiết tổng được, và phải đo tính tổng được bằng số như đã trình "
     "bày ở mục 7.6.", "thuong"),

    ("10.5. Điều rút ra chung", "de_muc_phu"),
    ("Bốn hướng trên thất bại vì bốn lý do khác nhau, và không hướng nào thất bại vì "
     "lỗi cài đặt. "
     "Hướng thứ nhất và hướng thứ hai đúng về lý thuyết nhưng không có chỗ phát huy "
     "trong chế độ vận hành thật. Hướng thứ ba sai ngay ở giả thiết, và được phát hiện "
     "bằng lập luận chứ không phải bằng đo. Hướng thứ tư đúng cả về lý thuyết lẫn ý "
     "tưởng nhưng vướng một tính chất của chính tập ràng buộc. Ghi lại cả bốn, kèm số "
     "liệu, là cách giữ cho các hướng này thật sự đóng chứ không mở lại theo trí nhớ. "
     "Mục 11 trình bày mã nguồn và cách tái lập, trong đó có cả các phép đo âm tính nói "
     "trên.", "thuong"),
]


# Bảng ở mục 11: các tệp mã nguồn chính và vai trò, số dòng đếm từ kho lưu trữ.
BANG_MUC_11 = {
    "tieu_de": ("Bảng 11.1. Các tệp mã nguồn chính của phần lý thuyết và thực nghiệm. "
                "Số dòng không tính dòng trắng."),
    "cot": ["Tệp", "Vai trò", "Số dòng"],
    "dong": [
        ("pie_net/constraints.py",
         "Phép chiếu lên quả cầu biến phân toàn phần, chứng chỉ sai số, phép ép khả thi",
         "377"),
        ("pie_net/reflected_solver.py",
         "Sơ đồ phản xạ hai dòng và các chế độ chọn ngân sách bước nội", "345"),
        ("pie_net/operators.py", "Toán tử làm mờ và ước lượng hằng số Lipschitz", "160"),
        ("pie_net/data.py", "Sinh dữ liệu thử nghiệm và tạo ảnh quan sát", "106"),
        ("grid_schedule.py", "Lưới so chi phí theo giao thức đo công bằng ở mục 8",
         "150"),
        ("test_certificate.py", "Kiểm chặn trên của chứng chỉ và đo cái giá của nó",
         "108"),
        ("test_summability.py", "Đo độ dốc và tính tổng được của dãy sai số", "92"),
        ("phan_tich_muc7.py", "Rút số liệu của mục 7 từ tệp kết quả", "155"),
        ("tests/test_theory.py", "Mười bốn kiểm thử tính chất toán học", "160"),
    ],
}

MUC_11 = [
    ("Mục 11. Chất lượng mã nguồn và tính tái lập", "de_muc"),

    ("Các con số ở mục 7 và các kết luận âm tính ở mục 10 chỉ có giá trị nếu người khác "
     "kiểm lại được. Mục này trình bày những việc đã làm để phần thực thi đủ tin cậy "
     "cho việc đó: bộ kiểm thử, việc dọn các khẳng định vượt quá số liệu, cách tổ chức "
     "mã nguồn, và cách chạy lại từng thí nghiệm.", "thuong"),

    ("11.1. Bộ kiểm thử", "de_muc_phu"),
    ("Có mười bốn kiểm thử tự động, tất cả đều chạy qua. Nguyên tắc khi viết chúng là "
     "kiểm tính chất toán học chứ không kiểm chi tiết cài đặt, vì một kiểm thử ràng "
     "vào chi tiết cài đặt sẽ hỏng mỗi lần sửa mã mà không cho biết điều gì về tính "
     "đúng đắn.", "thuong"),

    ("Vài kiểm thử tiêu biểu. Một kiểm thử xác nhận đặc trưng biến phân của phép chiếu, "
     "tức tích vô hướng giữa vectơ từ điểm chiếu tới điểm cần chiếu với hướng đi từ "
     "điểm chiếu tới mọi điểm khác của tập ràng buộc đều không dương. Một kiểm thử xác "
     "nhận chứng chỉ là chặn trên thật sự, tức nó không bao giờ nhỏ hơn sai số thật; "
     "đây chính là tính chất mà mục 6.5 đo bằng số, nay được đưa vào bộ kiểm thử để "
     "không thể vô tình làm hỏng về sau. Một kiểm thử xác nhận khởi tạo ấm tốn ít bước "
     "nội hơn khởi tạo lạnh. Một kiểm thử xác nhận độ dài bước sinh ra từ hằng số "
     "Lipschitz ước lượng được luôn thỏa điều kiện ở giả thiết thứ ba của mục 9.2. "
     "Các kiểm thử còn lại phủ tính khả thi của đầu ra, tính không âm của khoảng cách "
     "đối ngẫu, và quan hệ giữa ngân sách bước nội với sai số thu được.", "thuong"),

    ("11.2. Dọn các khẳng định vượt quá số liệu", "de_muc_phu"),
    ("Hai việc dọn dẹp đã thực hiện, cả hai đều là sửa lỗi của chính đề tài.", "thuong"),

    ("Việc thứ nhất là gỡ khẳng định rằng phần dư biến phân tiến về không. Khẳng định "
     "này từng có trong cả tài liệu hướng dẫn lẫn chú thích mã nguồn. Số liệu thật chỉ "
     "cho thấy phần dư biến phân giảm từ 2,48 xuống 0,199 qua 200 bước ngoài, và giảm "
     "không đơn điệu. Một dãy giảm không đơn điệu trên một đoạn hữu hạn không phải bằng "
     "chứng cho việc tiến về không, nên khẳng định cũ đã được thay bằng đúng điều số "
     "liệu cho phép nói.", "thuong"),

    ("Việc thứ hai là sửa danh pháp, đã trình bày ở mục 5.1. Pha thứ ba của sơ đồ bốn "
     "pha cũ là hiệu chỉnh kiểu Tseng chứ không phải bước phản xạ, và các tài liệu "
     "trước gọi sai tên. Tên gọi sai này đã được sửa trong tài liệu hướng dẫn và trong "
     "chú thích mã nguồn, kèm một dòng ghi rõ tên cũ sai ở chỗ nào, để người đọc tài "
     "liệu cũ không bị nhầm tiếp.", "thuong"),

    ("11.3. Tổ chức mã nguồn", "de_muc_phu"),
    ("Phần lý thuyết và thực nghiệm gồm chín tệp chính, tổng cộng 1653 dòng lệnh không "
     "tính dòng trắng. Bảng 11.1 liệt kê từng tệp cùng vai trò của nó.", "thuong"),

    ("", "bang11"),

    ("Cách chia tệp theo một nguyên tắc: mỗi tệp trả lời một câu hỏi. Ba tệp đo, gồm "
     "tệp kiểm chứng chỉ, tệp đo tính tổng được và tệp rút số liệu của mục 7, chạy độc "
     "lập với nhau và mỗi tệp ghi thẳng kết quả ra tệp dữ liệu. Nhờ vậy mọi con số "
     "trong báo cáo đều truy được về một tệp cụ thể thay vì về trí nhớ.", "thuong"),

    ("11.4. Tính tái lập", "de_muc_phu"),
    ("Mã nguồn, dữ liệu thô và các kịch bản phân tích đều công khai trên kho lưu trữ "
     "của đề tài. Dữ liệu thô gồm các tệp kết quả của mọi lần chạy, kể cả các lần chạy "
     "cho kết quả âm tính ở mục 10, chứ không chỉ các lần chạy cho kết quả thuận lợi. "
     "Lệnh chạy lại từng thí nghiệm được ghi trong phần chú thích đầu của mỗi tệp, kèm "
     "các tham số đã dùng, nên người kiểm không phải đoán.", "thuong"),

    ("Ngoài ra có một trình kiểm tra tự động cho chính báo cáo này, gồm bảy nhóm kiểm "
     "tra. Nó đối chiếu từng con số trong báo cáo với tệp kết quả, kể cả các số trong "
     "bảng và các số ghi ở dạng khoa học; kiểm tính nhất quán của thuật ngữ; kiểm cách "
     "viết hoa; kiểm các mục có dẫn sang nhau hay không; đếm lại số dòng lệnh của từng "
     "tệp trong bảng 11.1 để bảng đó không lỗi thời khi mã nguồn thay đổi; và kiểm mọi "
     "đường dẫn nêu trong bảng 12.1 có còn tồn tại hay không. Trình kiểm tra này đã bắt "
     "được nhiều lỗi trong quá trình viết, trong đó có những con số lấy từ một lần chạy "
     "cũ đã bị thay thế, và một đường dẫn ghi thiếu tên thư mục.", "thuong"),

    ("11.5. Hạ tầng tính toán", "de_muc_phu"),
    ("Các phép đo chi phí ở mục 7 chạy trên một máy chủ có bộ xử lý đồ họa, vì phép đo "
     "thời gian chỉ có nghĩa khi mọi cấu hình chạy trên cùng một máy. Có ba trở ngại về "
     "môi trường đã phải xử lý, ghi lại ở đây vì người khác dựng lại môi trường sẽ gặp "
     "đúng chúng.", "thuong"),

    ("Trở ngại thứ nhất là máy chủ không có sẵn công cụ tạo môi trường ảo của Python. "
     "Trở ngại thứ hai là không có sẵn trình quản lý gói, phải cài riêng vào thư mục "
     "người dùng. Trở ngại thứ ba, và mất thời gian nhất, là bản thư viện tính toán "
     "được tải về mặc định biên dịch cho một phiên bản nền tảng tính toán song song mới "
     "hơn phiên bản mà trình điều khiển của máy chủ hỗ trợ, nên phải chỉ định rõ nguồn "
     "tải bản khớp với trình điều khiển. Ba trở ngại này không ảnh hưởng tới kết quả, "
     "nhưng chúng giải thích vì sao phần thiết lập tốn thời gian hơn phần chạy.",
     "thuong"),

    ("Mục 12 tiếp theo liệt kê đầy đủ các sản phẩm đã bàn giao, kèm vị trí của từng "
     "sản phẩm trong kho lưu trữ.", "thuong"),
]


# Bảng ở mục 12: các sản phẩm bàn giao, vị trí đã đối chiếu với kho lưu trữ.
BANG_MUC_12 = {
    "tieu_de": ("Bảng 12.1. Các sản phẩm bàn giao của giai đoạn nghiên cứu. Vị trí "
                "tính từ thư mục gốc của kho lưu trữ."),
    "cot": ["Sản phẩm", "Vị trí", "Mô tả"],
    "dong": [
        ("Bản thảo bài báo bằng tiếng Anh",
         "tai_lieu_bai_bao/paper/",
         "Tệp nguồn và bản đã biên dịch, 8 trang, biên dịch không lỗi"),
        ("Tài liệu định vị",
         "tai_lieu_bai_bao/00_ket_luan_dieu_hanh.md và "
         "tai_lieu_bai_bao/01_khung_bai_bao.md",
         "Kết luận điều hành và khung bài báo, dùng để quyết định hướng đi"),
        ("Khảo sát tài liệu",
         "tai_lieu_bai_bao/khao_sat_tai_lieu_bang_tong_hop.md và "
         "tai_lieu_bai_bao/khao_sat_tai_lieu_ket_luan_5_goc.md",
         "Bảng thuộc tính của các bài đã đọc, và kết luận theo từng góc khảo sát"),
        ("Báo cáo số liệu",
         "tai_lieu_bai_bao/02_ket_qua_so_tien_hoa.md",
         "Toàn bộ số liệu đã đo, kể cả số liệu của các lần chạy đã bị thay thế"),
        ("Bản thảo chứng minh thứ nhất",
         "tai_lieu_bai_bao/03_chung_minh_dinh_ly.md",
         "Bản nhắm hội tụ mạnh cho sơ đồ bốn pha, đã bị bác, giữ kèm biên bản"),
        ("Lộ trình chứng minh",
         "tai_lieu_bai_bao/04_lo_trinh_chung_minh.md",
         "Các bước phải làm sau khi bản thảo thứ nhất bị bác"),
        ("Bản thảo chứng minh thứ hai",
         "tai_lieu_bai_bao/05_chung_minh_hoi_tu_yeu.md",
         "Bốn bổ đề đứng vững, phần còn lại bị bác, kèm biên bản phản biện"),
        ("Bản chứng minh hiện hành",
         "tai_lieu_bai_bao/06_chung_minh_day_du.md",
         "Bản đang dùng, gồm dẫn xuất hằng số nhiễu và kết quả phản biện"),
        ("Tệp kết quả thô",
         "results/theory/",
         "32 tệp số liệu và 4 tệp nhật ký chạy, gồm cả các lần chạy âm tính"),
        ("Mã nguồn và bộ kiểm thử",
         "pie_net/, tests/ và các tệp đo ở thư mục gốc",
         "1653 dòng lệnh và mười bốn kiểm thử, xem bảng 11.1"),
        ("Báo cáo này",
         "tai_lieu_bai_bao/bao_cao/",
         "Tệp báo cáo, trình sinh báo cáo và trình kiểm tra tự động sáu nhóm"),
    ],
}

MUC_12 = [
    ("Mục 12. Sản phẩm bàn giao", "de_muc"),

    ("Mục này liệt kê những gì đã bàn giao, để người hướng dẫn biết trong tay có sẵn "
     "cái gì và tìm chúng ở đâu. Bảng 12.1 ghi tên sản phẩm, vị trí trong kho lưu trữ, "
     "và một dòng mô tả. Mọi vị trí trong bảng đã được đối chiếu với kho lưu trữ tại "
     "thời điểm viết báo cáo.", "thuong"),

    ("", "bang12"),

    ("Có một điểm về cách bàn giao cần nói rõ. Hai bản thảo chứng minh đã bị bác vẫn "
     "được giữ lại nguyên trạng, kèm biên bản phản biện chỉ ra chỗ sai, thay vì xóa đi "
     "và chỉ nộp bản hiện hành. Lý do là hai bản ấy cho thấy vì sao sơ đồ cuối cùng có "
     "dạng như ở mục 9.1: bản thứ nhất sụp vì bước quán tính và bước neo cắt đứt chuỗi "
     "phép chiếu, và chính điều đó dẫn tới việc bỏ hai bước ấy ở mục 5. Người đọc chỉ "
     "xem bản hiện hành sẽ không thấy được lý do đó.", "thuong"),

    ("Tương tự, tệp kết quả thô giữ cả số liệu của các lần chạy đã bị thay thế và các "
     "lần chạy cho kết quả âm tính ở mục 10, chứ không chỉ giữ các lần chạy được trích "
     "vào báo cáo. Nhờ vậy người kiểm có thể tự đối chiếu con số trong báo cáo với "
     "nguồn, và cũng thấy được những con số đã bị loại cùng lý do loại.", "thuong"),

    ("Mục 13 tiếp theo nêu các hạn chế của công việc này.", "thuong"),
]


MUC_13 = [
    ("Mục 13. Hạn chế", "de_muc"),

    ("Mục này nêu năm hạn chế của công việc. Mục đích là để người đọc hiểu đúng phạm vi "
     "của các kết luận ở mục 7 và mục 9, chứ không phải để bào chữa. Mỗi hạn chế được "
     "nêu thẳng, kèm điều nó ngăn không cho kết luận.", "thuong"),

    ("13.1. Phạm vi thực nghiệm hẹp", "de_muc_phu"),
    ("Toàn bộ thực nghiệm chỉ dùng một tập ràng buộc là quả cầu biến phân toàn phần, "
     "và hai loại mờ là mờ Gauss và mờ chuyển động, trên ảnh xám cạnh 96 điểm ảnh với "
     "tám ảnh kiểm tra. Chứng chỉ ở mục 6 được dẫn cho riêng tập ràng buộc này, vì công "
     "thức của hàm mục tiêu đối ngẫu phụ thuộc dạng cụ thể của nó.", "thuong"),

    ("Điều này ngăn không cho kết luận rằng chế độ ngân sách thích nghi tiết kiệm được "
     "chi phí trên các tập ràng buộc khác. Cách gỡ vòng luẩn quẩn ở mục 6.2 chỉ cần bộ "
     "giải nội thuộc loại gốc và đối ngẫu, nên về nguyên tắc nó áp được cho tập ràng "
     "buộc khác; nhưng đó là suy đoán, chưa phải điều đã đo.", "thuong"),

    ("13.2. Chưa so sánh với phương pháp đã công bố", "de_muc_phu"),
    ("Đây là hạn chế nặng nhất. Mọi so sánh ở mục 7 đều là giữa các chế độ chọn ngân "
     "sách bước nội của cùng một sơ đồ: chế độ thích nghi so với phép chiếu chính xác "
     "và với hai tiêu chuẩn tương đối. Không có phương pháp nào đã công bố được cài đặt "
     "lại để so.", "thuong"),

    ("Hệ quả là các con số ở bảng 7.1 và bảng 7.2 chỉ nói được một điều: trong cùng một "
     "sơ đồ, chọn ngân sách bước nội theo chứng chỉ rẻ hơn chiếu chính xác. Chúng không "
     "nói được rằng sơ đồ này rẻ hơn các phương pháp khác cho cùng bài toán. Muốn nói "
     "được điều thứ hai thì phải cài lại ít nhất một phương pháp trong bảng 4.1 và đo "
     "trên cùng máy, việc đó chưa làm.", "thuong"),

    ("13.3. Bán kính quả cầu đặt theo thông tin của ảnh sạch", "de_muc_phu"),
    ("Trong mọi thực nghiệm, bán kính quả cầu biến phân toàn phần được đặt bằng 0,55 "
     "lần biến phân toàn phần của ảnh gốc chưa bị làm mờ. Đây là thông tin mà một "
     "triển khai thật không có, vì nếu đã biết ảnh gốc thì không cần khôi phục.",
     "thuong"),

    ("Cách đặt này được chọn có chủ ý, để bán kính giống nhau cho mọi cấu hình và phép "
     "so chi phí không bị nhiễu bởi việc chọn bán kính. Nó hợp lý cho mục đích so chi "
     "phí, nhưng phải nói rõ rằng nó làm bài toán dễ hơn thực tế: trong triển khai "
     "thật, bán kính phải ước lượng từ dữ liệu quan sát, và sai số của ước lượng đó sẽ "
     "cộng thêm vào sai số cuối cùng. Báo cáo không đo ảnh hưởng của việc ước lượng "
     "bán kính.", "thuong"),

    ("13.4. Lợi thế đo được là chi phí, không phải chất lượng ảnh", "de_muc_phu"),
    ("Cần nói rõ chế độ ngân sách thích nghi không cho ảnh đẹp hơn. Ở cấu hình tốt nhất "
     "của mỗi nhóm, chỉ số chất lượng ảnh của chế độ thích nghi là 26,8251 dB trên mờ "
     "Gauss so với 26,7998 dB của phép chiếu chính xác, và 26,8487 dB trên mờ chuyển "
     "động so với 26,8332 dB. Chênh lệch lần lượt là 0,03 dB và 0,02 dB, tức nằm trong "
     "khoảng dao động giữa các cấu hình và không có ý nghĩa thực tế.", "thuong"),

    ("Đây đúng là điều được kỳ vọng, vì hai chế độ giải cùng một bài toán và chỉ khác "
     "nhau ở cách phân bổ chi phí cho vòng lặp nội. Nhưng nó giới hạn cách phát biểu "
     "đóng góp: đóng góp là đạt cùng độ chính xác với chi phí thấp hơn, chứ không phải "
     "đạt độ chính xác cao hơn. Người đọc kỳ vọng ảnh khôi phục tốt hơn sẽ không tìm "
     "thấy điều đó ở đây.", "thuong"),

    ("13.5. Hai chi tiết trong chứng minh chưa viết đầy đủ", "de_muc_phu"),
    ("Như đã nêu ở mục 9.7, phần lý thuyết còn hai chỗ chưa hoàn tất: phát biểu chính "
     "xác của bổ đề tựa Fejér cho dạng truy hồi có thừa số nhân, và bước chuyển qua "
     "giới hạn yếu trong không gian vô hạn chiều. Trong không gian hữu hạn chiều của "
     "phần thực nghiệm, cả hai là thường quy, nên hạn chế này không ảnh hưởng tới các "
     "con số đã báo.", "thuong"),

    ("Tuy vậy nó ảnh hưởng tới việc nộp bài. Chừng nào hai chỗ ấy chưa viết đủ thì định "
     "lý chưa nộp được, và bài chỉ còn phần thuật toán cùng phần thực nghiệm. Đây là "
     "việc phải làm trước, và mục 14 xếp nó vào nhóm ưu tiên cao nhất.", "thuong"),

    ("Năm hạn chế trên không làm hỏng các kết luận đã nêu, nhưng chúng giới hạn phạm vi "
     "của các kết luận đó. Mục 14 trình bày việc còn lại, sắp theo thứ tự ưu tiên rút "
     "ra từ chính năm hạn chế này.", "thuong"),
]


MUC_14 = [
    ("Mục 14. Việc còn lại và khuyến nghị", "de_muc"),

    ("Mục này liệt kê ba việc còn lại, sắp theo thứ tự ưu tiên rút ra từ năm hạn chế ở "
     "mục 13, kèm ước lượng công sức cho từng việc. Cuối mục là một khuyến nghị về mức "
     "công bố. Cả ba việc đều đã đủ rõ để quyết định ngay, không cần khảo sát thêm.",
     "thuong"),

    ("14.1. Việc thứ nhất: hoàn tất hai chi tiết của chứng minh", "de_muc_phu"),
    ("Đây là việc ưu tiên cao nhất, vì chừng nào chưa xong thì phần định lý chưa nộp "
     "được. Hai chi tiết đã nêu ở mục 9.7 và nhắc lại ở mục 13.5: phát biểu chính xác "
     "của bổ đề tựa Fejér cho dạng truy hồi có thừa số nhân, và bước chuyển qua giới "
     "hạn yếu trong không gian vô hạn chiều.", "thuong"),

    ("Cả hai đều là việc tra cứu và viết cho đủ, không phải việc tìm ý mới, vì hướng "
     "xử lý đã biết và chỉ cần dùng đúng phát biểu có sẵn trong tài liệu. Ước lượng "
     "công sức: một đến hai tuần.", "thuong"),

    ("14.2. Việc thứ hai: bổ sung so sánh với phương pháp đã công bố", "de_muc_phu"),
    ("Đây là việc ưu tiên thứ hai, và lý do rất thực tế: người phản biện chắc chắn sẽ "
     "hỏi. Như mục 13.2 đã nêu, mọi so sánh hiện có đều nằm trong cùng một sơ đồ, nên "
     "chưa trả lời được câu hỏi sơ đồ này đứng ở đâu so với các phương pháp khác.",
     "thuong"),

    ("Đề xuất cụ thể là cài lại hai hoặc ba phương pháp trong bảng 4.1 và đo trên cùng "
     "máy, theo đúng giao thức đo ở mục 8. Nên ưu tiên phương pháp chiếu phản xạ gốc và "
     "một phương pháp dùng phép chiếu xấp xỉ với tiêu chuẩn sai số tương đối, vì đó là "
     "hai bài gần công việc này nhất. Ước lượng công sức: ba đến năm tuần, phần lớn dành "
     "cho việc cài lại cho đúng chứ không phải cho việc chạy.", "thuong"),

    ("14.3. Việc thứ ba: chọn tạp chí", "de_muc_phu"),
    ("Ba tạp chí phù hợp với mức đóng góp đã định vị là Optimization, Journal of "
     "Computational and Applied Mathematics, và Numerical Algorithms. Đề xuất lấy "
     "Optimization làm lựa chọn chính, vì phạm vi của nó nhận cả phần thuật toán lẫn "
     "phần thực nghiệm, và vì mức đóng góp của bài khớp với mặt bằng bài đăng ở đó.",
     "thuong"),

    ("Cần nói rõ một thay đổi so với ban đầu. Ở giai đoạn khảo sát tài liệu, tạp chí "
     "được nhắm tới là Numerical Algorithms, và đó là lý do bài đăng năm 2023 trên tạp "
     "chí này được kiểm rủi ro trùng lặp trước mọi việc khác, như mục 4.4 đã trình bày. "
     "Sau khi biết phần định lý chỉ là mở rộng, thứ tự ưu tiên được đổi lại như trên; "
     "Numerical Algorithms lùi xuống phương án dự phòng chứ không bị loại. Quyết định "
     "cuối cùng thuộc về người hướng dẫn; phần này chỉ nêu căn cứ.", "thuong"),

    ("14.4. Khuyến nghị về mức công bố", "de_muc_phu"),
    ("Khuyến nghị là không nhắm nhóm Q1 với phần định lý hiện tại. Đây không phải sự "
     "thận trọng quá mức mà là kết luận rút từ chính phép tính ở mục 9.6: hằng số của "
     "số hạng nhiễu điểm gốc không chứa nghịch đảo độ dài bước, nên việc thêm bước phản "
     "xạ vào khung phép chiếu xấp xỉ không tạo ra cơ chế mới, và định lý là một mở "
     "rộng.", "thuong"),

    ("Các tạp chí nhóm hàng đầu như SIAM Journal on Optimization hay Mathematical "
     "Programming đòi một cơ chế mới. Nộp vào đó với phần định lý hiện tại nhiều khả "
     "năng chỉ dẫn tới một vòng phản biện dài rồi bị từ chối, và thời gian ấy đủ để làm "
     "xong cả hai việc ở mục 14.1 và 14.2.", "thuong"),

    ("Cách đặt giá trị của bài nên theo đúng những gì đã đo được: đóng góp nằm ở chứng "
     "chỉ sai số tính được, ở chế độ ngân sách thích nghi cùng phân tích chi phí, và ở "
     "giao thức đo công bằng. Ba thứ đó thuộc về thuật toán và thực nghiệm, và chúng "
     "không yếu đi vì phần định lý chỉ là mở rộng.", "thuong"),

    ("Mục 15 là phụ lục, gồm các bảng số liệu đầy đủ và các chi tiết kỹ thuật không đưa "
     "vào phần chính.", "thuong"),
]


BANG_PL_A = {
    "tieu_de": ("Bảng 15.1. Phụ lục A: bảng thuật ngữ. Cột tiếng Việt là từ được dùng "
                "thống nhất trong toàn báo cáo."),
    "cot": ["Tiếng Việt", "Tiếng Anh"],
    "dong": [
        ("bất đẳng thức biến phân", "variational inequality"),
        ("tập ràng buộc", "constraint set"),
        ("tập nghiệm", "solution set"),
        ("toán tử chi phí", "cost operator"),
        ("quả cầu biến phân toàn phần", "total variation ball"),
        ("điểm phản xạ", "reflected point"),
        ("bước quán tính", "inertial step"),
        ("bước neo", "anchor step"),
        ("hiệu chỉnh kiểu Tseng", "Tseng-type correction"),
        ("phép chiếu xấp xỉ", "inexact projection"),
        ("bước nội", "inner iteration"),
        ("bước ngoài", "outer iteration"),
        ("khởi tạo ấm", "warm start"),
        ("chứng chỉ sai số", "error certificate"),
        ("khoảng cách đối ngẫu", "duality gap"),
        ("ngân sách thích nghi", "adaptive budget"),
        ("tổng được", "summable"),
        ("phần dư biến phân", "variational residual"),
        ("đại lượng thế năng", "potential function"),
        ("tựa Fejér", "quasi-Fejér"),
        ("độ dài bước", "step size"),
    ],
}

BANG_PL_B = {
    "tieu_de": ("Bảng 15.2. Phụ lục B: lệnh chạy lại từng thí nghiệm. Chạy từ thư mục "
                "gốc của kho lưu trữ."),
    "cot": ["Thí nghiệm", "Lệnh", "Mục liên quan"],
    "dong": [
        ("So chi phí theo giao thức đo công bằng",
         "python grid_schedule.py --device cuda --K 150 --size 96 --n_test 8 "
         "--blur gauss", "mục 7, mục 8"),
        ("Kiểm chặn trên của chứng chỉ và đo cái giá",
         "python test_certificate.py", "mục 6"),
        ("Đo tính tổng được của dãy sai số",
         "python test_summability.py --device cuda --K 600", "mục 7.6"),
        ("Rút số liệu của mục 7 từ tệp kết quả",
         "python phan_tich_muc7.py", "mục 7"),
        ("Chạy bộ kiểm thử", "python -m pytest tests/", "mục 11.1"),
        ("Sinh lại báo cáo này",
         "python tai_lieu_bai_bao/bao_cao/tao_bao_cao.py", "mục 11.4"),
        ("Kiểm tra báo cáo",
         "python tai_lieu_bai_bao/bao_cao/kiem_tra_bao_cao.py", "mục 11.4"),
    ],
}

BANG_PL_C = {
    "tieu_de": ("Bảng 15.3. Phụ lục C: danh mục tệp kết quả, nhóm theo thí nghiệm sinh "
                "ra chúng. Tất cả nằm trong thư mục results/theory/."),
    "cot": ["Nhóm tệp", "Số tệp", "Thuộc thí nghiệm nào"],
    "dong": [
        ("certificate_check.csv, certificate_cost.csv", "2",
         "Kiểm chứng chỉ, nguồn của bảng 6.1 và của các số ở mục 6.6"),
        ("grid_fair_gauss.csv, grid_fair_motion.csv và hai tệp pareto", "4",
         "Lưới so chi phí công bằng, nguồn của bảng 7.1 và bảng 7.2"),
        ("grid_schedule_gauss.csv, grid_schedule_motion.csv", "2",
         "Lưới lịch đầy đủ, nguồn của các số ở mục 8.1 và mục 10.1"),
        ("muc7_he_so_*.csv, muc7_kha_thi.csv, muc7_lich_*.csv, muc7_do_nhay.csv", "6",
         "Số liệu rút cho mục 7, sinh bởi phan_tich_muc7.py"),
        ("summability_gauss.csv, summability_motion.csv", "2",
         "Đo độ dốc và tính tổng được, nguồn của mục 7.6"),
        ("theory_gauss_*.csv, theory_motion_*.csv", "12",
         "Quỹ đạo theo từng chế độ ngân sách, nguồn của mức khả thi ở mục 7.4"),
        ("pseudomono_*.csv, summary.csv", "4",
         "Thí nghiệm về toán tử giả đơn điệu, không trích vào báo cáo này"),
        ("compare_accel_gauss.log và ba tệp nhật ký khác", "4",
         "Nhật ký chạy, nguồn của bảng 8.1 và của các số ở mục 10.1"),
        ("results/theory/phan_tich.md", "1", "Ghi chép phân tích kèm theo các lần chạy"),
    ],
}

BANG_PL_D = {
    "tieu_de": ("Bảng 15.4. Phụ lục D: biên bản các vòng phản biện. Kết quả của mỗi "
                "vòng đều đã được tích hợp vào báo cáo."),
    "cot": ["Vòng", "Đối tượng phản biện", "Kết quả"],
    "dong": [
        ("1", "Bản thảo chứng minh thứ nhất, nhắm hội tụ mạnh cho sơ đồ bốn pha",
         "Bị bác. Bổ đề một bước đòi chuỗi phép chiếu liền mạch, mà bước quán tính và "
         "bước neo cắt đứt chuỗi đó. Dẫn tới việc bỏ hai bước ấy, xem mục 5"),
        ("2", "Số liệu so chi phí bản đầu",
         "Bị bác. Ba lỗi đo lường được chỉ ra, dẫn tới giao thức đo công bằng ở mục 8"),
        ("3", "Bản thảo chứng minh thứ hai, hội tụ yếu",
         "Bốn bổ đề đứng vững, phần lập luận tính mới bị bác vì là ngụy biện. Đề xuất "
         "thay bằng tiêu chuẩn chiếu nới lỏng, hướng này về sau đóng lại, xem mục 10.4"),
        ("4", "Bước ghép đại lượng thế năng cho phép chiếu chính xác",
         "Ba người phản biện độc lập tính lại từng phép biến đổi và đều xác nhận đúng, "
         "kể cả đẳng thức then chốt, xem mục 9.4"),
        ("5", "Khẳng định hằng số nhiễu chứa nghịch đảo độ dài bước",
         "Bị bác. Hai người phản biện xác nhận hằng số không chứa nghịch đảo độ dài "
         "bước, nên định lý là mở rộng, xem mục 9.6"),
    ],
}

MUC_15 = [
    ("Mục 15. Phụ lục", "de_muc"),

    ("Phụ lục gồm bốn phần. Phụ lục A là bảng thuật ngữ, để người đọc đối chiếu với "
     "tài liệu tiếng Anh và để việc dùng từ trong báo cáo nhất quán. Phụ lục B là các "
     "lệnh chạy lại thí nghiệm. Phụ lục C là danh mục tệp kết quả kèm chú thích từng "
     "nhóm thuộc thí nghiệm nào. Phụ lục D là biên bản các vòng phản biện.", "thuong"),

    ("15.1. Phụ lục A: bảng thuật ngữ", "de_muc_phu"),
    ("Trong toàn báo cáo, mỗi khái niệm chỉ dùng đúng một từ tiếng Việt. Việc này được "
     "kiểm tự động bằng nhóm kiểm tra thứ hai của trình kiểm tra nói ở mục 11.4, với "
     "danh sách các cặp từ đồng nghĩa không được dùng lẫn. Bảng 15.1 liệt kê các thuật "
     "ngữ chính.", "thuong"),

    ("", "bangA"),

    ("15.2. Phụ lục B: lệnh chạy lại thí nghiệm", "de_muc_phu"),
    ("Bảng 15.2 ghi lệnh chạy lại từng thí nghiệm, kèm mục của báo cáo mà thí nghiệm "
     "đó sinh số liệu. Lệnh so chi phí phải chạy hai lần, một lần với mờ Gauss và một "
     "lần với mờ chuyển động, bằng cách đổi giá trị của tham số cuối. Các lệnh có tham "
     "số chỉ định bộ xử lý đồ họa sẽ chạy được trên bộ xử lý trung tâm nếu đổi tham số "
     "đó, nhưng khi ấy con số thời gian không so được với con số trong báo cáo.",
     "thuong"),

    ("", "bangB"),

    ("15.3. Phụ lục C: danh mục tệp kết quả", "de_muc_phu"),
    ("Thư mục kết quả có 37 tệp, gồm 32 tệp số liệu, 4 tệp nhật ký chạy và một tệp ghi "
     "chép phân tích. Bảng 15.3 nhóm chúng theo thí nghiệm sinh ra chúng, để người kiểm "
     "biết tệp nào ứng với con số nào trong báo cáo.", "thuong"),

    ("", "bangC"),

    ("Một nhóm tệp được giữ lại dù không trích vào báo cáo, là nhóm thí nghiệm về toán "
     "tử giả đơn điệu. Lý do giữ đã nêu ở mục 12: dữ liệu thô giữ cả những lần chạy "
     "không dùng tới, để người kiểm thấy được toàn bộ những gì đã đo chứ không chỉ phần "
     "được trích.", "thuong"),

    ("15.4. Phụ lục D: biên bản các vòng phản biện", "de_muc_phu"),
    ("Có năm vòng phản biện trong giai đoạn nghiên cứu này. Bốn vòng kết thúc bằng việc "
     "bác một phần hoặc toàn bộ nội dung được xem xét, và cả bốn lần kết luận của người "
     "phản biện đều đúng; vòng còn lại xác nhận nội dung được xem xét là đúng. Bảng "
     "15.4 tóm tắt từng vòng.", "thuong"),

    ("", "bangD"),

    ("Các vòng phản biện này là lý do vì sao báo cáo có mục 8 và mục 10. Hai mục ấy ghi "
     "lại chỗ báo cáo tự nêu sai sót của chính mình, và chúng được giữ nguyên độ dài "
     "thay vì rút gọn, vì chính chúng là căn cứ để tin các con số còn lại.", "thuong"),

    ("Đến đây báo cáo kết thúc. Toàn bộ nội dung, từ mục 1 đến mục 15, đã được trình "
     "kiểm tra tự động soát qua bảy nhóm kiểm tra trước khi nộp.", "thuong"),
]


def thay_chu(p, text):
    """Thay chữ của một đoạn nhưng giữ nguyên định dạng của lần chạy chữ đầu tiên."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def xoa_doan(p):
    p._element.getparent().remove(p._element)


def ke_vien(t):
    """Kẻ viền cho bảng bằng XML, vì tệp khuôn không có sẵn kiểu bảng kẻ viền."""
    tbl_pr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for canh in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{canh}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tbl_pr.append(borders)


def dat_chu(o, text, dam=False, co=12):
    """Ghi chữ vào một ô của bảng, dùng phông Times New Roman."""
    o.text = ""
    p = o.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(co)
    r.bold = dam
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15


def them_bang(doc, bang):
    """Thêm một bảng hai cột kèm dòng tiêu đề bảng ở trên."""
    p = doc.add_paragraph()
    r = p.add_run(bang["tieu_de"])
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    t = doc.add_table(rows=1, cols=len(bang["cot"]))
    ke_vien(t)
    for o, ten in zip(t.rows[0].cells, bang["cot"]):
        dat_chu(o, ten, dam=True)
    for dong in bang["dong"]:
        o = t.add_row().cells
        for i, gia_tri in enumerate(dong):
            dat_chu(o[i], gia_tri)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def main():
    if not os.path.exists(KHUON):
        sys.exit(f"Không tìm thấy tệp khuôn: {KHUON}")
    doc = Document(KHUON)

    # 1. Thay chữ trên bìa, giữ nguyên hai ảnh khung viền và logo.
    for i, text in BIA.items():
        thay_chu(doc.paragraphs[i], text)

    # 2. Xóa toàn bộ nội dung cũ phía sau ngắt trang, gồm cả bảng.
    for p in list(doc.paragraphs[DOAN_NGAT_TRANG + 1:]):
        xoa_doan(p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # 3. Ghi các mục nội dung.
    bang_theo_khoa = {
        "bang": BANG_MUC_4, "bang6": BANG_MUC_6, "bang7a": BANG_MUC_7A,
        "bang7b": BANG_MUC_7B, "bang8": BANG_MUC_8, "bang11": BANG_MUC_11,
        "bang12": BANG_MUC_12,
        "bangA": BANG_PL_A, "bangB": BANG_PL_B,
        "bangC": BANG_PL_C, "bangD": BANG_PL_D,
    }
    for text, kieu in (MUC_1 + MUC_2 + MUC_3 + MUC_4 + MUC_5 + MUC_6 + MUC_7
                       + MUC_8 + MUC_9 + MUC_10 + MUC_11 + MUC_12 + MUC_13
                       + MUC_14 + MUC_15):
        if kieu in bang_theo_khoa:
            them_bang(doc, bang_theo_khoa[kieu])
            continue
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        if kieu == "de_muc":
            r.bold = True
            r.font.size = Pt(15)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(12)
        elif kieu == "de_muc_phu":
            r.bold = True
            r.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        else:
            r.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Pt(28)
        p.paragraph_format.line_spacing = 1.4

    doc.save(OUT)
    print(f"đã tạo: {OUT}")


if __name__ == "__main__":
    main()
